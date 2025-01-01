"""UNIT_test_summary_report.py: Used to generate summary report based on excel ( RA, etc. ) as input
Parameters has been given from argparse command line interface.
As input, it is needed an path where template docx is
As output, it is needed a path to generate an docx file based upon given template
"""

__author__ = 'Ardeleanu Lucian'
__copyright__ = "Copyright 2021, Autoliv Electronic"
__version__ = "1.0"
__email__ = 'lucian.ardeleanu@autoliv.com'
__status__ = 'Released'

# ============ IMPORTS ====================
import datetime
import argparse
import sys
import os
import glob
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Pt
from docx.shared import RGBColor

import openpyxl

currentDir = os.getcwd()
sys.path.append(currentDir)


# ===============================================================
# MAIN FUNCTION OF SCRIPT
# INPUT ARGUMENT: NONE
# OUTPUT ARGUMENT: NONE
# ===============================================================
def main():

    parser = argparse.ArgumentParser()
    parser.add_argument('-i', '--input-path-template-docx', help="Path to Template DOCX ( Summary Report ) with filename and extension", required=True)
    parser.add_argument('-o', '--output-path-generated-docx', help="Path to generate DOCX Summary Report with filename and extension",required=True)
    args = parser.parse_args()

    path_to_template_docx_file = args.input_path_template_docx
    path_to_generated_docx_file = args.output_path_generated_docx

    # -i="Audi_Tranche_6_Unit_Tests_Summary_Report_Template.docx" -o="Generated.docx"
    # path_to_template_docx_file = "Audi_Tranche_6_Unit_Tests_Summary_Report_Template.docx"
    # path_to_generated_docx_file = "Generated.docx"

    path_to_unit_test_reports = "S:\Components\Application"
    print("Please make sure to have S: drive mapped!")

    # Initialise an empty dict
    dict_with_all_data_from_unit_test_reports = {}

    # For every file that meet conditions in path
    for root, dirs, files in os.walk(path_to_unit_test_reports):
        for file in files:
            if (file.endswith(".xlsx") and '_UnitTestReport' in  file):

                # Get full path to unit test excel file
                path_to_file = os.path.join(root, file)

                # Open file
                wb_obj = openpyxl.load_workbook(path_to_file)

                # For every sheet in all sheets
                for sheet_name in wb_obj.sheetnames:

                    # If sheet is not Summary, so open every sheet except summarry
                    if sheet_name != "Summary" and "Template" not in sheet_name:

                        # Open sheet
                        sheet = wb_obj[sheet_name]

                        list_of_req_per_sheet = []

                        # Variable used to count rows
                        count_rows = 0
                        for row in sheet.iter_rows():

                            if row[0].value != None  and row[0].value != "" and not "Test Case" in row[0].value and not "Test_Case" in row[0].value:

                                list_per_row = ['', '', '']

                                # append to list TC ( Test Case )
                                if row[0].value != None and row[0].value != '' and row[0].value != ' ' :
                                    list_per_row[0] = row[0].value

                                # append to list Result of test
                                if row[3].value != None and row[0].value != '' and row[0].value != ' ' :
                                    list_per_row[1] = row[3].value

                                # append to list Requirements
                                if row[4].value != row[0].value and row[4].value != None:
                                    list_per_row[2] = row[4].value.replace('\n', " ")
                                else:
                                    # If requirement not found on same row as test case and result
                                    if row[0].value != None and row[0].value != '' and row[0].value != ' ' :

                                        # For Debug purpose, print sheet and test case with no requirement
                                        # print(sheet_name, row[0].value, row[4].value)

                                        # Range of rows to search
                                        for search_row_count in range(0,3):
                                            # if a Requirement has been found in any range of rows and Requirement column, then keep it
                                            if "_" in str(sheet[count_rows + search_row_count][4].value):
                                                list_per_row[2] = sheet[count_rows + search_row_count][4].value.replace('\n', " ")

                                        if list_per_row[2] == '':
                                            print("Test case " + row[0].value + " found in sheet " + sheet_name + " in file " + file + " has no requirement!" )

                                if ( list_per_row[0] != list_per_row[1] != list_per_row[2] and list_per_row not in list_of_req_per_sheet ):
                                    list_of_req_per_sheet.append(list_per_row)
                            # Count number of rows
                            count_rows += 1

                        dict_with_all_data_from_unit_test_reports[sheet_name] = list_of_req_per_sheet

    if len(dict_with_all_data_from_unit_test_reports) == 0:
        print('DATA CANNOT BE EXTRACTED FROM GIVEN INPUT, NO DOCUMENT WILL BE GENERATED')
        return 0
    else:
        print('DATA HAS BEEN EXTRACTED SUCCESFULLY!')

    # For debug purpose
    # for key, value in dict_with_all_data_from_unit_test_reports.items():
    #     print( key, ' : ', value )

    # ====================================================================
    # WRITE EXTRACTED DATA IN DOCX DOCUMENT
    # ====================================================================

    # Get an instance of document
    document = Document(path_to_template_docx_file)
    document.save(path_to_generated_docx_file)

    # Get all values from dictionary in a list
    list_of_values_from_dictionary = list(dict_with_all_data_from_unit_test_reports.values())
    list_of_keys_from_dictionary = list(dict_with_all_data_from_unit_test_reports.keys())

    number_of_test_case = 0
    for sheet in list_of_values_from_dictionary:
        number_of_test_case += len(sheet)

    # Open document
    document = Document(path_to_generated_docx_file)

    for table in document.tables:

        # ==========================
        # Add / delete rows to table
        # ==========================

        difference_of_rows = 0
        # If found table is correct
        if table.cell(0, 0).text == "TC" and table.cell(0, 1).text == "Result" and table.cell(0, 2).text == "Requirements":
            # Variable used to count number of rows
            number_of_rows = 0

            for row in table.rows:
                # Get number of maximum rows
                number_of_rows += 1

            # Get difference of number of rows between excel list and table
            difference_of_rows = number_of_rows - ( number_of_test_case + len(list_of_keys_from_dictionary) + 1 )

            # If there are more rows than needed
            if (difference_of_rows > 0):
                # Delete exceeded rows
                for counter_for_difference_rows in range(0, difference_of_rows):
                    row = table.rows[difference_of_rows - counter_for_difference_rows]
                    remove_row(table, row)
            else:
                # Else, add rows to table
                for counter_for_difference_rows in range(0, -difference_of_rows):
                    table.add_row()

        # ==========================
        # Add data in table
        # ==========================

        # If found table is correct
        counter_for_rows = 1
        if table.cell(0, 0).text == "TC" and table.cell(0, 1).text == "Result" and table.cell(0,2).text == "Requirements":
            for key, value in dict_with_all_data_from_unit_test_reports.items():

                # ======== WRITE IN  CELLS ==========

                # Write cell that contains sheet name

                # ACCESS CELL FROM TABLE
                cell = table.cell(counter_for_rows, 0)  # first parameter is row, second is column

                # access paragraph from cell
                paragraph_from_cell = cell.paragraphs[0]

                # add text to it and bold it
                run = paragraph_from_cell.add_run(key)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)

                counter_for_rows += 1

                for test_case in value:

                    # WRITE FIRST CELL: TEST CASE

                    # ACCESS CELL FROM TABLE
                    cell = table.cell(counter_for_rows, 0)  # first parameter is row, second is column

                    # access paragraph from cell
                    paragraph_from_cell = cell.paragraphs[0]

                    # add text to it and bold it
                    run = paragraph_from_cell.add_run(test_case[0])
                    run.bold = True
                    run.font.color.rgb = RGBColor(48, 84, 150)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

                    # WRITE SECOND CELL: RESULT

                    # ACCESS CELL FROM TABLE
                    cell = table.cell(counter_for_rows, 1)  # first parameter is row, second is column

                    # access paragraph from cell
                    paragraph_from_cell = cell.paragraphs[0]

                    # add text to it and bold it
                    run = paragraph_from_cell.add_run(test_case[1])
                    run.bold = False
                    run.font.color.rgb = RGBColor(48, 84, 150)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

                    # WRITE THIRD CELL: REQUIREMENT

                    # ACCESS CELL FROM TABLE
                    cell = table.cell(counter_for_rows, 2)  # first parameter is row, second is column

                    # access paragraph from cell
                    paragraph_from_cell = cell.paragraphs[0]

                    # add text to it and bold it
                    run = paragraph_from_cell.add_run(test_case[2])
                    run.bold = False
                    run.font.color.rgb = RGBColor(48, 84, 150)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

                    counter_for_rows += 1

    # List all paragraphs in document
    paragraphs = list(document.paragraphs)

    # Iterare and index all paragraphs
    for idx, paragraph in enumerate(paragraphs):
        # if found a specific paragraph
        if 'Number of tests:' in paragraph.text:
            run = paragraph.add_run( str(number_of_test_case) )
            run.bold = True
            run.italic = True

        if 'Number of tests passed:' in paragraph.text:
            run = paragraph.add_run( str(number_of_test_case) )
            run.bold = True
            run.italic = True

    document.save(path_to_generated_docx_file)

if __name__ == '__main__':
    main()

