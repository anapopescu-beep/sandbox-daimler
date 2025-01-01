"""qac_summary_report.py: Used to generate a docx file based on xlsx inputs
Parameters has been given from argparse command line interface.
As input, it is needed 2 paths where QAC_Metrics_Report and QAC_Warning_Report are
As output, it is needed a path to generate an docx file
"""

__author__ = 'Ardeleanu Lucian'
__copyright__ = "Copyright 2021, Autoliv Electronic"
__version__ = "1.0"
__email__ = 'lucian.ardeleanu@autoliv.com'
__status__ = 'Released'

# ================ IMPORTS ==================
import openpyxl
from docx import Document
import argparse

parser = argparse.ArgumentParser()
parser.add_argument('-i', '--input-path-metrics-report',help="Path to XLSX QAC_Metrics_Report", required=True)
parser.add_argument('-l', '--input-path-warning-report',help="Path to XLSX QAC_Warning_Report", required=True)
parser.add_argument('-m', '--input-path-docx-report',help="Path to DOCX QAC_Metrics_Analysis_Report", required=True)
args = parser.parse_args()

path_to_QAC_Metrics_Report_excel = args.input_path_metrics_report
path_to_QAC_Warning_Report_excel = args.input_path_warning_report
path_to_QAC_Metrics_Analisys_Report_docx = args.input_path_docx_report


# ===================
# GLOBAL FUNCTIONS
# ===================

# ===============================================================
# FUNCTION USED TO REMOVE A ROW FROM A GIVEN TABLE
# INPUT ARGUMENT: OBJECT-BASED TABLE, OBJECT-BASED ROW
# OUTPUT ARGUMENT: NONE
# ===============================================================
def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


# ====================================================================
# EXTRACT DATA FROM QAC METRICS REPORT EXCEL FILE
# ====================================================================
# Open workbook
wb = openpyxl.load_workbook(path_to_QAC_Metrics_Report_excel)
sheets = wb.sheetnames

# open sheet
# Variable to check if code checklist sheet is found
sheet_found_flag = 0
for sheet_name in sheets:

    if ('Improvable' in sheet_name.lower() or 'improvable' == sheet_name.lower()):
        ws = wb[sheet_name]
        sheet_found_flag = 1
        break
    else:
        sheet_found_flag = 0
if sheet_found_flag == 0:
    print('Improvable SHEET NOT FOUND on document ' + path_to_QAC_Metrics_Report_excel)
else:
    # Configure rows and columns to read exactly what is needed
    rows = ws.max_row

    #columns = ws.max_column
    columns = 6

    # append document data in list

    list_of_data_from_QAC_Metrics_Report_excel = []
    for i in range(3, rows + 1):
        data_from_row = []
        for j in range(1, columns):
            data_from_row.append(ws.cell(row=i, column=j).value)
        list_of_data_from_QAC_Metrics_Report_excel.append(data_from_row)

# ====================================================================
# END OF EXTRACT DATA FROM QAC METRICS REPORT EXCEL FILE
# ====================================================================


# ====================================================================
# EXTRACT DATA FROM QAC WARNING REPORT EXCEL FILE
# ====================================================================
# Open workbook
wb = openpyxl.load_workbook(path_to_QAC_Warning_Report_excel)
sheets = wb.sheetnames

# open sheet
# Variable to check if code checklist sheet is found
sheet_found_flag = 0
for sheet_name in sheets:

    if ('metrics' in sheet_name.lower() or 'Metrics' == sheet_name.lower()):
        ws = wb[sheet_name]
        sheet_found_flag = 1
        break
    else:
        sheet_found_flag = 0
if sheet_found_flag == 0:
    print('metrics SHEET NOT FOUND on document ' + path_to_QAC_Warning_Report_excel)
else:
    # Configure rows and columns to read exactly what is needed
    rows = ws.max_row - 1

    #columns = ws.max_column
    columns = 7

    # append document data in list

    list_of_data_from_QAC_Warning_Report_excel = []
    for i in range(6, rows):
        data_from_row = []
        for j in range(1, columns):
            data_from_row.append(ws.cell(row=i, column=j).value)
        list_of_data_from_QAC_Warning_Report_excel.append(data_from_row)

# ====================================================================
# END OF EXTRACT DATA FROM QAC WARNING REPORT EXCEL FILE
# ====================================================================


# ====================================================================
# WRITE EXTRACTED DATA IN DOCX DOCUMENT
# ====================================================================

# Open document
document = Document(path_to_QAC_Metrics_Analisys_Report_docx)

for table in document.tables:

    # If found table is correct
    if table.cell(0, 0).text =="File Name" and table.cell(0, 1).text =="Function Name" :
        # Variable used to count number of rows
        number_of_rows = 0

        for row in table.rows:
            # Get number of maximum rows
            number_of_rows += 1

        # Get difference of number of rows between excel list and table
        difference_of_rows = number_of_rows - len(list_of_data_from_QAC_Metrics_Report_excel)

        # If there are more rows than needed
        if ( difference_of_rows > 0):
            # Delete exceeded rows
            for counter_for_difference_rows in range(0, difference_of_rows ):
                row = table.rows[difference_of_rows - counter_for_difference_rows ]
                remove_row(table, row)
        else:
            # Else, add rows to table
            for counter_for_difference_rows in range(0, -difference_of_rows):
                table.add_row()

        # Do not allow table to autofit itself
        table.autofit = False
        table.allow_autofit = False

    if table.cell(0, 0).text =="Component" and table.cell(0, 1).text =="Warnings " :
        # Variable used to count number of rows
        number_of_rows = 0

        for row in table.rows:
            # Get number of maximum rows
            number_of_rows += 1

        # Get difference of number of rows between excel list and table
        difference_of_rows = number_of_rows - len(list_of_data_from_QAC_Warning_Report_excel)

        # If there are more rows than needed
        if ( difference_of_rows > 0):

            # Delete exceeded rows
            for counter_for_difference_rows in range(0, difference_of_rows ):
                row = table.rows[difference_of_rows - counter_for_difference_rows ]
                remove_row(table, row)
        else:
            # Else, add rows to table
            for counter_for_difference_rows in range(0, -difference_of_rows):
                table.add_row()


# Write data in table extracted from excel file
for table in document.tables:
    # Write metrics report in docx
    if table.cell(0, 0).text =="File Name" and table.cell(0, 1).text =="Function Name" :
        for row in range(1, len(list_of_data_from_QAC_Metrics_Report_excel)):
            for column in range(0, len(list_of_data_from_QAC_Metrics_Report_excel[0])):
                table.cell(row ,column ).text = str(list_of_data_from_QAC_Metrics_Report_excel[row][column])

    # Write warning report in docx
    if table.cell(0, 0).text =="Component" and table.cell(0, 1).text =="Warnings " :
        for row in range(1, len(list_of_data_from_QAC_Warning_Report_excel)):
            for column in range(0, len(list_of_data_from_QAC_Warning_Report_excel[0])):
                table.cell(row ,column ).text = str(list_of_data_from_QAC_Warning_Report_excel[row][column])

document.save(path_to_QAC_Metrics_Analisys_Report_docx)

# ====================================================================
# END OF WRITE EXTRACTED DATA IN DOCX DOCUMENT
# ====================================================================