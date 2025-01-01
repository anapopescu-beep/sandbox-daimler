"""Qualification_test_summary_report.py: Used to generate summary report based on excel ( RA, etc. ) as input
Parameters has been given from argparse command line interface.
As input, it is needed an path where tests can be found and to template
As output, it is needed a path to generate an docx file based upon given template
"""

__author__ = 'Ardeleanu Lucian'
__copyright__ = "Copyright 2021, Autoliv Electronic"
__version__ = "1.0"
__email__ = 'lucian.ardeleanu@autoliv.com'
__status__ = 'Released'


import xlrd
import xlrd2
import openpyxl
import os
from docx import Document
from copy import deepcopy
from docx.shared import Pt
from docx.shared import Cm, Inches
try:
    import win32com.client
except:
    pass
import argparse

# ===============================================================
# FUNCTION USED TO UPDATE THE TABLE OF CONTENT FOR DOCX FILES
# INPUT ARGUMENT: PATH TO DOCUMENT
# OUTPUT ARGUMENT: NONE
# ===============================================================
def update_toc(docx_file):
    word = win32com.client.DispatchEx("Word.Application")
    doc = word.Documents.Open(docx_file)
    toc_count = doc.TablesOfContents.Count
    if toc_count != 0:
        doc.TablesOfContents(1).Update()
        print('TOC should have been updated.')
    else:
        print('TOC has not been updated for sure...')
    doc.Close(SaveChanges=True)
    word.Quit()


# ===============================================================
# FUNCTION USED TO REMOVE A ROW FROM A GIVEN TABLE
# INPUT ARGUMENT: OBJECT-BASED TABLE, OBJECT-BASED ROW
# OUTPUT ARGUMENT: NONE
# ===============================================================
def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


# ===============================================================
# FUNCTION USED TO INSERT PARAGRAPH AFTER SPECIFIC PARAGRAPHS
# INPUT ARGUMENT: LIST OF OBJECT-BASED PARAGRAPHS, INTEGER: INDEX,
#                 STRING: TEXT TO INSERT
# OUTPUT ARGUMENT: IT WILL INSERT A PARAGRAPH INSIDE DOCUMENT
# ===============================================================
def insert_paragraph_after(paragraphs, idx, text=None):
    next_paragraph_idx = idx + 1
    if idx == len(paragraphs):
        return document.add_paragraph(text)
    next_paragraph = paragraphs[next_paragraph_idx]
    return next_paragraph.insert_paragraph_before(text)


# ===============================================================
# FUNCTION USED TO OPEN EXCEL FILE AND EXTRACT DATA FROM IT
# INPUT ARGUMENT: STRING: PATH TO EXCEL FILE TO PROCESS
# OUTPUT ARGUMENT: LIST OF STRINGS WITH DATA FROM EXCEL
# ===============================================================
def extract_data_from_xls(file_to_process):

    # Define List in witch data will be stored
    raw_list = []

    # open workbook
    wb = xlrd.open_workbook(file_to_process)
    sheets = wb.sheet_names()

    # open sheet
    # check which file is needed for opening

    for sheet_name in sheets:
        if ( 'SRM' in sheet_name.lower() or 'srm' in sheet_name.lower() ):
            ws = wb[sheet_name]
            break

    # Get total number of rows and columns
    rows = ws.nrows

    # We don't actually need all columns for document
    #columns = ws.ncols
    columns = 7

    # append document data in list
    for i in range(3, rows):
        for j in range(6, columns):

            # Extract cell value in variable
            cell_value = ws.cell_value(i,j)

            # If cell value is not empty
            if cell_value != '':

                # Remove all spaces
                cell_value = cell_value.replace(" ", "")

                # Replace , with ;
                if ',' in cell_value:
                    cell_value = cell_value.replace("," , ";")

                # Split after every ; and append in list
                if ";" in cell_value:
                    for element in cell_value.split(';'):
                        if element not in raw_list:
                            raw_list.append(element)
                else:
                    if cell_value not in raw_list:
                        raw_list.append(cell_value)
    return raw_list


# ===============================================================
# FUNCTION USED TO OPEN EXCEL FILE AND EXTRACT DATA FROM IT
# INPUT ARGUMENT: STRING: PATH TO EXCEL FILE TO PROCESS
# OUTPUT ARGUMENT: LIST OF STRINGS WITH DATA FROM EXCEL
# ===============================================================
def extract_data_from_xlsm_review_covearge(file_to_process):

    # open workbook
    wb = xlrd2.open_workbook(file_to_process)

    # open sheet
    sheets = wb.sheets()

    # check which file is needed for opening
    for sheet in sheets:
        if ( 'Tests Synthesis' in sheet.name.lower() or 'tests synthesis' in sheet.name.lower() ):
            ws = sheet
            break

    # Get total number of rows and columns
    rows = ws.nrows

    # append document data in list
    list_of_global_data_from_excel =[]

    # Initialise first list that will contain only tests from xls
    list_with_test_cases = []

    for rows_counter in range(2, rows):

        # extract cell value
        cell_value = ws.cell_value(rows_counter, 2)

        # if cell is a test case
        list_with_data_per_row = ['', '', '']
        if 'QUALIF' in cell_value:
            # Extract test case data ( TEST CASE, STATUS, ISSUE )
            list_with_data_per_row[0] =  ws.cell_value(rows_counter, 2)
            list_with_data_per_row[1] =  ws.cell_value(rows_counter, 10).upper()
            list_with_data_per_row[2] =  str(ws.cell_value(rows_counter, 12)).split('.')[0]

            # append test case in global list
            if ( list_with_data_per_row != ['', '', ''] ):
                list_with_test_cases.append(list_with_data_per_row)

        if list_with_test_cases not in list_of_global_data_from_excel and list_with_test_cases != []:
            # the list_of_global_data_from_excel contain [ test category, list of test cases]
            # append test category
            list_of_global_data_from_excel.append(file_to_process.split('Peer_Review\\')[1].split('_Review_Coverage')[0])
            # append test cases list
            list_of_global_data_from_excel.append(list_with_test_cases)

    return list_of_global_data_from_excel


# ===============================================================
# FUNCTION USED TO OPEN EXCEL FILE AND EXTRACT DATA FROM IT
# INPUT ARGUMENT: STRING: PATH TO EXCEL FILE TO PROCESS
# OUTPUT ARGUMENT: LIST OF STRINGS WITH DATA FROM EXCEL
# ===============================================================
def extract_data_from_xlsm_manual_tests(file_to_process):

    # open workbook
    wb = xlrd2.open_workbook(file_to_process)

    # open sheet
    sheets = wb.sheets()

    # check which file is needed for opening
    for sheet in sheets:
        if ( 'Tests Synthesis' in sheet.name.lower() or 'tests synthesis' in sheet.name.lower() ):
            ws = sheet
            break

    # Get total number of rows and columns
    rows = ws.nrows

    # append document data in list
    list_of_global_data_from_excel =[]

    # Initialise first list that will contain only tests from xls
    list_with_test_cases = []

    for rows_counter in range(2, rows):

        # extract cell value
        cell_value = ws.cell_value(rows_counter, 5)

        # if cell is a test case
        list_with_data_per_row = ['', '', '']
        if 'QUALIF' in cell_value:
            # Extract test case data ( TEST CASE, STATUS, ISSUE )
            list_with_data_per_row[0] =  ws.cell_value(rows_counter, 5)
            list_with_data_per_row[1] =  ws.cell_value(rows_counter, 10).upper()
            list_with_data_per_row[2] =  str(ws.cell_value(rows_counter, 12)).split('.')[0]

            # append test case in global list
            if ( list_with_data_per_row != ['', '', ''] ):
                list_with_test_cases.append(list_with_data_per_row)

    return list_with_test_cases


# ===============================================================
# FUNCTION USED TO OPEN XML FILES AND EXTRACT TESTS STATUS
# INPUT ARGUMENT:STRING: PATH TO XML FILE TO PROCESS
# OUTPUT ARGUMENT:STRING: STATUS OF THE TEST
# ===============================================================
def extract_data_from_xml(file_to_process):
    # status is read flag
    status_is_read = 0
    # initialize status as OK
    status_return = "NOK"
    # Open file for reading text.
    with open(file_to_process, 'r') as myfile:
        # For each line of text
        for line in myfile:
            # verify if the line contain "Result Test:" and if the line not contain "Passed".
            if "Result Test:" in line:
                status_is_read = 1
                if 'PASSED' in line:
                    status_return = "OK"
            if status_is_read:
                break
    return status_return


# ===============================================================
# MAIN FUNCTION OF SCRIPT
# INPUT ARGUMENT: NONE
# OUTPUT ARGUMENT: NONE
# ===============================================================
def main():

    parser = argparse.ArgumentParser()
    parser.add_argument('-i', '--input-path-test-qualification', help="Path to Qualification Tests folder", required=True)
    parser.add_argument('-l', '--input-path-template-docx', help="Path to Template DOCX ( Summary Report )  with filename and extension", required=True)
    parser.add_argument('-o', '--output-path-generated-docx', help="Path to generate DOCX Summary Report with filename and extension",required=True)
    args = parser.parse_args()

    path_to_test_qualification = args.input_path_test_qualification
    path_to_template_docx_file = args.input_path_template_docx
    path_to_generated_docx_file = args.output_path_generated_docx

    path_to_xlsm_peer_review_folder = path_to_test_qualification + '\\Peer_Review'

    # Initialize test counters
    number_of_OK_manual_tests = []
    number_of_NOK_manual_tests = []
    number_of_total_OK_manual_tests = 0
    number_of_total_NOK_manual_tests = 0
    number_of_OK_manual_tests.append(0)
    number_of_NOK_manual_tests.append(0)

    number_of_OK_automatic_tests = []
    number_of_NOK_automatic_tests = []
    number_of_total_OK_automatic_tests = 0
    number_of_total_NOK_automatic_tests = 0
    number_of_OK_automatic_tests.append(0)
    number_of_NOK_automatic_tests.append(0)

    number_of_OK_tests_review_coverage = []
    number_of_NOK_tests_review_coverage = []
    number_of_total_OK_tests_review_coverage = 0
    number_of_total_NOK_tests_review_coverage = 0
    number_of_OK_tests_review_coverage.append(0)
    number_of_NOK_tests_review_coverage.append(0)

    # ====================================================================
    # Get a dict of all automatic tests
    # ====================================================================
    dict_with_all_automatic_test_scripts = {}
    for root, dirs, files in os.walk(path_to_test_qualification + '\\Automatic_Tests\\Test_Scenarios'):
        if root.endswith('Test_Scripts'):

            list_of_tests = []

            # extract module name from file name
            module_name = root.split('\\')[len(root.split('\\')) - 2]

            for file in files:
                if file.endswith(".ts") :

                    test_name = file.replace('.ts', '')

                    # read the status of test from .xml files
                    xml_path = path_to_test_qualification + '\\Automatic_Tests\\Test_Reports\\' + module_name + '\\' + file.replace('.ts', '.xml')
                    try:
                        test_status = extract_data_from_xml(xml_path)
                    except:
                        test_status = 'NOTRUN'

                    # list of tests and status
                    list_of_tests.append([test_name, test_status])

            # append data in dictionary
            dict_with_all_automatic_test_scripts[module_name] = list_of_tests

    # Check if dict is empty or not
    if len(dict_with_all_automatic_test_scripts) == 0:
        print('DATA CANNOT BE EXTRACTED FROM AUTOMATIC TESTS, NO DOCUMENT WILL BE GENERATED')
        return 0
    else:
        print('DATA HAS BEEN EXTRACTED SUCCESSFULLY FROM AUTOMATIC TESTS!')


    # ====================================================================
    # Get a dict of all manual tests
    # ====================================================================
    dict_with_all_manual_test_scripts = {}
    for root, dirs, files in os.walk(path_to_test_qualification + '\\Manual_tests\\Test_Scenarios'):
        if len(files) > 0 and root.split('\\')[len(root.split('\\')) - 2] == 'Test_Scenarios':
            list_of_tests_for_this_module = []
            # extract module name from file name
            module_name = root.split('\\')[len(root.split('\\')) - 1]
            for file in files:
                if file.endswith(".xlsm"):
                    list_of_tests_for_this_module = extract_data_from_xlsm_manual_tests(root + "\\" + file)

            # append data in dictionary
            dict_with_all_manual_test_scripts[module_name] = list_of_tests_for_this_module

    for root, dirs, files in os.walk((path_to_test_qualification + '\\Manual_tests\\Test_Reports')):
        if len(files) > 0 and root.split('\\')[len(root.split('\\')) - 2] == 'Test_Reports':
            list_of_tests_for_this_module = []
            # extract module name from file name
            module_name = root.split('\\')[len(root.split('\\')) - 1]
            for file in files:
                if file.endswith(".xml"):

                    test_name = file.replace('.xml', '')
                    # read the status of test from .xml file
                    test_status = extract_data_from_xml(root + '\\' + file)

                    # list of tests and status
                    list_of_tests_for_this_module.append([test_name, test_status, ''])

            if module_name in dict_with_all_manual_test_scripts.keys():
                # add data in main list
                for elem1 in list_of_tests_for_this_module:
                    elem_in_dict_flag = False
                    for elem2 in dict_with_all_manual_test_scripts[module_name]:
                        if elem2[0] == elem1[0]:
                            elem_in_dict_flag = True
                    if not elem_in_dict_flag:
                        dict_with_all_manual_test_scripts[module_name].append(elem1)
            else:
                dict_with_all_manual_test_scripts[module_name] = list_of_tests_for_this_module


    # Check if dict is empty or not
    if len(dict_with_all_manual_test_scripts) == 0:
        print('DATA CANNOT BE EXTRACTED FROM MANUAL TESTS, NO DOCUMENT WILL BE GENERATED')
        return 0
    else:
        print('DATA HAS BEEN EXTRACTED SUCCESSFULLY FROM MANUAL TESTS!')


    # ====================================================================
    # Extract Review Coverage data from excel
    # ====================================================================
    dict_with_all_data_from_review_coverage = []
    for root, dirs, files in os.walk(path_to_xlsm_peer_review_folder):
        for file in files:
            if file.endswith("Review_Coverage.xlsm"):
                dict_with_all_data_from_review_coverage.append(extract_data_from_xlsm_review_covearge(root + '\\' + file))

    # Check if dict is empty or not
    if len(dict_with_all_manual_test_scripts) == 0:
        print('DATA CANNOT BE EXTRACTED FROM REVIEW COVERAGE, NO DOCUMENT WILL BE GENERATED')
        return 0
    else:
        print('DATA HAS BEEN EXTRACTED SUCCESSFULLY FROM REVIEW COVERAGE!')

    # ====================================================================
    # WRITE EXTRACTED DATA IN DOCX DOCUMENT
    # ====================================================================

    # Get an instance of document
    document = Document(path_to_template_docx_file)
    document.save(path_to_generated_docx_file)

    # Get all values from dictionary in a list
    list_of_automatic_tests = list(dict_with_all_automatic_test_scripts.values())
    list_of_module_names_for_automatic_tests = list(dict_with_all_automatic_test_scripts.keys())

    # Get all values from dictionary in a list
    list_of_manual_tests = list(dict_with_all_manual_test_scripts.values())
    list_of_module_names_for_manual_tests = list(dict_with_all_manual_test_scripts.keys())

    # ===================================
    # ADD HEADINGS AND TABLES TO DOCUMENT
    # ===================================

    # add headings and tables for manual tests
    for counter in range(0, len(list_of_module_names_for_manual_tests) - 1):
        # Open document
        document = Document(path_to_generated_docx_file)

        # Select table needed to be copied
        template = document.tables[4]

        # Here we do the copy of the table
        tbl = template._tbl
        new_tbl = deepcopy(tbl)

        # List all paragraphs in document
        paragraphs = list(document.paragraphs)

        # Iterare and index all white box tests paragraphs
        for idx, paragraph in enumerate(paragraphs):
            # if found a specific paragraph
            if 'Global Summary for manual tests' in  paragraph.text :
                # add Title
                prior_paragraph = paragraphs[idx - 1].insert_paragraph_before("Tests Performed – ")
                prior_paragraph.style = 'Heading 3'

                # add Table
                prior_paragraph = paragraph.insert_paragraph_before(paragraphs[idx - 1]._p.addnext(new_tbl))

        document.save(path_to_generated_docx_file)

    # add headings and tables for automatic tests
    for counter in range(0, len(list_of_module_names_for_automatic_tests) - 1):
        # Open document
        document = Document(path_to_generated_docx_file)

        # Select table needed to be copied
        template = document.tables[4]

        # Here we do the copy of the table
        tbl = template._tbl
        new_tbl = deepcopy(tbl)

        # List all paragraphs in document
        paragraphs = list(document.paragraphs)

        # Iterare and index all white box tests paragraphs
        for idx, paragraph in enumerate(paragraphs):
            # if found a specific paragraph
            if 'Global Summary for automatic tests' in paragraph.text:
                # add Title
                prior_paragraph = paragraphs[idx - 1].insert_paragraph_before("Tests Performed – ")
                prior_paragraph.style = 'Heading 3'

                # add Table
                prior_paragraph = paragraph.insert_paragraph_before(paragraphs[idx - 1]._p.addnext(new_tbl))

        document.save(path_to_generated_docx_file)

    # add headings and tables for Reviews requirements
    for counter in range(0, len(dict_with_all_data_from_review_coverage) - 1):
        # Open document
        document = Document(path_to_generated_docx_file)

        # Select table needed to be copied
        template = document.tables[4]

        # Here we do the copy of the table
        tbl = template._tbl
        new_tbl = deepcopy(tbl)

        # List all paragraphs in document
        paragraphs = list(document.paragraphs)

        # Iterare and index all white box tests paragraphs
        for idx, paragraph in enumerate(paragraphs):
            # if found a specific paragraph
            if 'Global Summary for Reviews requirements tests' in paragraph.text:
                # add Title
                prior_paragraph = paragraphs[idx - 1].insert_paragraph_before("Tests Performed – ")
                prior_paragraph.style = 'Heading 3'

                # add Table
                prior_paragraph = paragraph.insert_paragraph_before(paragraphs[idx - 1]._p.addnext(new_tbl))

        document.save(path_to_generated_docx_file)


    # =================================
    # ADD DATA IN INSERTED PARAGRAPHS
    # =================================

    # Open document
    document = Document(path_to_generated_docx_file)

    # List all paragraphs in document
    paragraphs = list(document.paragraphs)

    # counter_for_list_of_values_from_dict = 0
    counter_for_list_of_values_from_dict_manual_tests = 0
    counter_for_list_of_values_from_dict_automatic_tests = 0
    counter_for_list_of_values_from_dict_review_coverage = 0

    for table in document.tables[4: 4 + len(list_of_module_names_for_automatic_tests) + len(list_of_module_names_for_manual_tests) + len(dict_with_all_data_from_review_coverage)]:
        # ==============================
        # for manual tests
        # ==============================
        if counter_for_list_of_values_from_dict_manual_tests < len(list_of_module_names_for_manual_tests):
            # COUNT HOW MANY ROWS ARE IN TABLE
            number_of_table_rows = 0
            for row in table.rows:
                # Get number of maximum rows
                number_of_table_rows += 1

            # GET DIFFERENCE OF ROWS AND ADD IF NEEDED MORE
            difference_of_rows = (number_of_table_rows - 1) - len(
                list_of_manual_tests[counter_for_list_of_values_from_dict_manual_tests])

            # If there are more rows than needed=
            if (difference_of_rows > 0):
                # Delete exceeded rows
                for counter_for_difference_rows in range(0, difference_of_rows):
                    row = table.rows[difference_of_rows - counter_for_difference_rows]
                    remove_row(table, row)
            else:
                # Else, add rows to table
                for counter_for_difference_rows in range(0, -difference_of_rows):
                    table.add_row()

            # Do not allow table to autofit itself
            table.autofit = False
            table.allow_autofit = False
            for index_test in range(0, len(list_of_manual_tests[counter_for_list_of_values_from_dict_manual_tests])):

                # ======== WRITE IN FIRST CELL ==========
                # ACCESS CELL FROM TABLE
                cell = table.cell(index_test + 1, 0) # first parameter is row, second is column

                # access paragraph from cell
                paragraph_from_cell = cell.paragraphs[0]

                # add text to it and bold it
                run = paragraph_from_cell.add_run(list_of_manual_tests[counter_for_list_of_values_from_dict_manual_tests][index_test][0].upper())
                run.bold = True

                # ======== WRITE IN SECOND CELL ==========
                # ACCESS CELL FROM TABLE
                cell = table.cell(index_test + 1, 1) # first parameter is row, second is column

                # access paragraph from cell
                paragraph_from_cell = cell.paragraphs[0]

                # add text to it and bold it
                run = paragraph_from_cell.add_run("SBE")
                run.bold = True

                # ======== WRITE IN THIRD CELL ==========
                # ACCESS CELL FROM TABLE
                cell = table.cell(index_test + 1, 2) # first parameter is row, second is column

                # access paragraph from cell
                paragraph_from_cell = cell.paragraphs[0]

                # add text to it and bold it
                test_status = list_of_manual_tests[counter_for_list_of_values_from_dict_manual_tests][index_test][1]
                run = paragraph_from_cell.add_run(test_status)
                run.bold = True

                # ======== WRITE IN FOURTH CELL ==========
                # ACCESS CELL FROM TABLE
                cell = table.cell(index_test + 1, 3)  # first parameter is row, second is column

                # access paragraph from cell
                paragraph_from_cell = cell.paragraphs[0]

                # add issue tracking number
                issue = list_of_manual_tests[counter_for_list_of_values_from_dict_manual_tests][index_test][2]
                run = paragraph_from_cell.add_run(issue)
                run.bold = True

                # count OK and NOK tests in lists of counters
                if test_status == "OK":
                    number_of_OK_manual_tests[counter_for_list_of_values_from_dict_manual_tests] += 1
                if test_status == "NOK":
                    number_of_NOK_manual_tests[counter_for_list_of_values_from_dict_manual_tests] += 1

            # count the total number of OK and NOK tests
            number_of_total_OK_manual_tests += number_of_OK_manual_tests[counter_for_list_of_values_from_dict_manual_tests]
            number_of_total_NOK_manual_tests += number_of_NOK_manual_tests[counter_for_list_of_values_from_dict_manual_tests]

            counter_for_list_of_values_from_dict_manual_tests += 1

            # append new values in counter list
            number_of_OK_manual_tests.append(0)
            number_of_NOK_manual_tests.append(0)

            # document.save(path_to_generated_docx_file)

        # ==============================
        # for automatic tests
        # ==============================
        elif counter_for_list_of_values_from_dict_automatic_tests < len(list_of_module_names_for_automatic_tests):
            # COUNT HOW MANY ROWS ARE IN TABLE
            number_of_table_rows = 0
            for row in table.rows:
                # Get number of maximum rows
                number_of_table_rows += 1

            # GET DIFFERENCE OF ROWS AND ADD IF NEEDED MORE
            difference_of_rows = (number_of_table_rows - 1) - len(
                list_of_automatic_tests[counter_for_list_of_values_from_dict_automatic_tests])

            # If there are more rows than needed=
            if (difference_of_rows > 0):
                # Delete exceeded rows
                for counter_for_difference_rows in range(0, difference_of_rows):
                    row = table.rows[difference_of_rows - counter_for_difference_rows]
                    remove_row(table, row)
            else:
                # Else, add rows to table
                for counter_for_difference_rows in range(0, -difference_of_rows):
                    table.add_row()

            # Do not allow table to autofit itself
            table.autofit = False
            table.allow_autofit = False
            for index_test in range(0, len(list_of_automatic_tests[counter_for_list_of_values_from_dict_automatic_tests])):

                # ======== WRITE IN FIRST CELL ==========
                # ACCESS CELL FROM TABLE
                cell = table.cell(index_test + 1, 0) # first parameter is row, second is column

                # access paragraph from cell
                paragraph_from_cell = cell.paragraphs[0]

                # add text to it and bold it
                run = paragraph_from_cell.add_run(list_of_automatic_tests[counter_for_list_of_values_from_dict_automatic_tests][index_test][0].upper())
                run.bold = True

                # ======== WRITE IN SECOND CELL ==========
                # ACCESS CELL FROM TABLE
                cell = table.cell(index_test + 1, 1) # first parameter is row, second is column

                # access paragraph from cell
                paragraph_from_cell = cell.paragraphs[0]

                # add text to it and bold it
                run = paragraph_from_cell.add_run("SBE")
                run.bold = True

                # ======== WRITE IN THIRD CELL ==========
                # ACCESS CELL FROM TABLE
                cell = table.cell(index_test + 1, 2) # first parameter is row, second is column

                # access paragraph from cell
                paragraph_from_cell = cell.paragraphs[0]

                # add text to it and bold it
                test_status = list_of_automatic_tests[counter_for_list_of_values_from_dict_automatic_tests][index_test][1]
                run = paragraph_from_cell.add_run(test_status)
                run.bold = True

                # count OK and NOK tests in lists of counters
                if test_status == "OK":
                    number_of_OK_automatic_tests[counter_for_list_of_values_from_dict_automatic_tests] += 1
                if test_status == "NOK":
                    number_of_NOK_automatic_tests[counter_for_list_of_values_from_dict_automatic_tests] += 1

            # count the total number of OK and NOK tests
            number_of_total_OK_automatic_tests += number_of_OK_automatic_tests[counter_for_list_of_values_from_dict_automatic_tests]
            number_of_total_NOK_automatic_tests += number_of_NOK_automatic_tests[counter_for_list_of_values_from_dict_automatic_tests]

            counter_for_list_of_values_from_dict_automatic_tests += 1

            # append new values in counter list
            number_of_OK_automatic_tests.append(0)
            number_of_NOK_automatic_tests.append(0)

            # document.save(path_to_generated_docx_file)

        # ==============================
        # for review requirements tests
        # ==============================
        elif counter_for_list_of_values_from_dict_review_coverage < len(dict_with_all_data_from_review_coverage):
            # COUNT HOW MANY ROWS ARE IN TABLE
            number_of_table_rows = 0
            for row in table.rows:
                # Get number of maximum rows
                number_of_table_rows += 1

            # GET DIFFERENCE OF ROWS AND ADD IF NEEDED MORE
            difference_of_rows = (number_of_table_rows - 1) - len(
                dict_with_all_data_from_review_coverage[counter_for_list_of_values_from_dict_review_coverage][1])

            # If there are more rows than needed=
            if (difference_of_rows > 0):
                # Delete exceeded rows
                for counter_for_difference_rows in range(0, difference_of_rows):
                    row = table.rows[difference_of_rows - counter_for_difference_rows]
                    remove_row(table, row)
            else:
                # Else, add rows to table
                for counter_for_difference_rows in range(0, -difference_of_rows):
                    table.add_row()

            # Do not allow table to autofit itself
            table.autofit = False
            table.allow_autofit = False
            for index_test in range(0, len(dict_with_all_data_from_review_coverage[counter_for_list_of_values_from_dict_review_coverage][1]) ):

                # ======== WRITE IN FIRST CELL ==========
                # ACCESS CELL FROM TABLE
                cell = table.cell(index_test + 1, 0) # first parameter is row, second is column

                # access paragraph from cell
                paragraph_from_cell = cell.paragraphs[0]

                # add name of test case
                run = paragraph_from_cell.add_run( dict_with_all_data_from_review_coverage[counter_for_list_of_values_from_dict_review_coverage][1][index_test][0].upper() )
                run.bold = True

                # ======== WRITE IN SECOND CELL ==========
                # ACCESS CELL FROM TABLE
                cell = table.cell(index_test + 1, 1) # first parameter is row, second is column

                # access paragraph from cell
                paragraph_from_cell = cell.paragraphs[0]

                # add text to it and bold it
                run = paragraph_from_cell.add_run("SBE")
                run.bold = True

                # ======== WRITE IN THIRD CELL ==========
                # ACCESS CELL FROM TABLE
                cell = table.cell(index_test + 1, 2) # first parameter is row, second is column

                # access paragraph from cell
                paragraph_from_cell = cell.paragraphs[0]

                # add test status
                test_status = dict_with_all_data_from_review_coverage[counter_for_list_of_values_from_dict_review_coverage][1][index_test][1]
                run = paragraph_from_cell.add_run(test_status)
                run.bold = True

                # ======== WRITE IN FOURTH CELL ==========
                # ACCESS CELL FROM TABLE
                cell = table.cell(index_test + 1, 3)  # first parameter is row, second is column

                # access paragraph from cell
                paragraph_from_cell = cell.paragraphs[0]

                # add issue tracking number
                issue = dict_with_all_data_from_review_coverage[counter_for_list_of_values_from_dict_review_coverage][1][index_test][2]
                run = paragraph_from_cell.add_run(issue)
                run.bold = True

                # count OK and NOK tests in lists of counters
                if test_status == "OK":
                    number_of_OK_tests_review_coverage[counter_for_list_of_values_from_dict_review_coverage] += 1
                if test_status == "NOK":
                    number_of_NOK_tests_review_coverage[counter_for_list_of_values_from_dict_review_coverage] += 1

            # count the total number of OK and NOK tests
            number_of_total_OK_tests_review_coverage = number_of_total_OK_tests_review_coverage + \
                                                       number_of_OK_tests_review_coverage[counter_for_list_of_values_from_dict_review_coverage]
            number_of_total_NOK_tests_review_coverage = number_of_total_NOK_tests_review_coverage +\
                                                        number_of_NOK_tests_review_coverage[counter_for_list_of_values_from_dict_review_coverage]

            counter_for_list_of_values_from_dict_review_coverage += 1

            # append new values in counter list
            number_of_OK_tests_review_coverage.append(0)
            number_of_NOK_tests_review_coverage.append(0)

    # document.save(path_to_generated_docx_file)
    counter_for_keys_from_dict_manual_tests = 0
    counter_for_keys_from_dict_automatic_tests = 0
    counter_for_keys_from_dict_review_coverage = 0
    # Add modules names to modules tested list
    for idx, paragraph in enumerate(paragraphs):
        # if found a specific paragraph

        if 'Tests Performed –' in paragraph.text:
            # Add dictionary key ( module names ) to previously added headings

            # add test cases in headers for manual tests
            if counter_for_keys_from_dict_manual_tests < len(list_of_module_names_for_manual_tests):
                run = paragraph.add_run(list_of_module_names_for_automatic_tests[counter_for_keys_from_dict_manual_tests])
                counter_for_keys_from_dict_manual_tests += 1

            # add test cases in headers for automatic tests
            elif counter_for_keys_from_dict_automatic_tests < len(list_of_module_names_for_automatic_tests):
                run = paragraph.add_run(list_of_module_names_for_automatic_tests[counter_for_keys_from_dict_automatic_tests])
                counter_for_keys_from_dict_automatic_tests += 1

            # add test cases in headers for review requirements tests
            elif counter_for_keys_from_dict_review_coverage < len(dict_with_all_data_from_review_coverage):
                run = paragraph.add_run(dict_with_all_data_from_review_coverage[counter_for_keys_from_dict_review_coverage][0])
                counter_for_keys_from_dict_review_coverage += 1

        if counter_for_keys_from_dict_manual_tests + counter_for_keys_from_dict_automatic_tests + counter_for_keys_from_dict_review_coverage >= \
                len(list_of_module_names_for_manual_tests) + len(list_of_module_names_for_automatic_tests) + len(dict_with_all_data_from_review_coverage):
            break


    # Add modules names to modules tested list
    for idx, paragraph in enumerate(paragraphs):
        # if found a specific paragraph
        if 'Priority #1 requirement tested:' in paragraph.text:
            # add an empty paragraph after upper paragraph
            prior_paragraph = insert_paragraph_after(paragraphs, idx, "\n")
            prior_paragraph.style = 'Normal'
            for key, value in dict_with_all_automatic_test_scripts.items():
                run = prior_paragraph.add_run(" - " + key + "\n")
            break

    # Count in list of values from dictionary
    counter_for_list_of_values_from_dictionary_manual_tests = 0
    counter_for_list_of_values_from_dictionary_automatic_tests = 0
    counter_for_list_of_values_from_dictionary_review_coverage = 0

    counter_of_all_tests_manual_tests = 0
    counter_of_all_tests_automatic_tests = 0
    counter_of_all_tests_review_coverage = 0

    # Iterare all paragraphs and add statistics for each module
    for idx, paragraph in enumerate(paragraphs):
        # for Manual tests
        if counter_for_list_of_values_from_dictionary_manual_tests < len(list_of_manual_tests) and 'Tests Performed' in paragraph.text:
            # Add an empty paragraph
            prior_paragraph = insert_paragraph_after(paragraphs, idx + 1, "\r\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run('Number of tests: ' + str(
                len(list_of_manual_tests[counter_for_list_of_values_from_dictionary_manual_tests])))
            run.bold = True
            run.italic = True

            # Add a new line to document
            prior_paragraph.add_run("\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run(
                'Number of positives tests: ' + str(number_of_OK_manual_tests[counter_for_list_of_values_from_dictionary_manual_tests]))
            run.bold = True
            run.italic = True

            # Add a new line to document
            prior_paragraph.add_run("\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run(
                'Number of negatives tests: ' + str(number_of_NOK_manual_tests[counter_for_list_of_values_from_dictionary_manual_tests]))
            run.bold = True
            run.italic = True

            # Add a new line to document
            prior_paragraph.add_run("\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run('Number of not run tests: '
                                          + str(len(list_of_manual_tests[counter_for_list_of_values_from_dictionary_manual_tests])
                                                - number_of_OK_manual_tests[counter_for_list_of_values_from_dictionary_manual_tests]
                                                - number_of_NOK_manual_tests[counter_for_list_of_values_from_dictionary_manual_tests]))
            run.bold = True
            run.italic = True

            # count all the white box tests
            counter_of_all_tests_manual_tests += len(list_of_automatic_tests[counter_for_list_of_values_from_dictionary_manual_tests])
            counter_for_list_of_values_from_dictionary_manual_tests += 1

        # for Automatic tests
        elif counter_for_list_of_values_from_dictionary_automatic_tests < len(list_of_automatic_tests) and 'Tests Performed' in paragraph.text:
            # Add an empty paragraph
            prior_paragraph = insert_paragraph_after(paragraphs, idx + 1, "\r\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run('Number of tests: ' + str(
                len(list_of_automatic_tests[counter_for_list_of_values_from_dictionary_automatic_tests])))
            run.bold = True
            run.italic = True

            # Add a new line to document
            prior_paragraph.add_run("\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run(
                'Number of positives tests: ' + str(number_of_OK_automatic_tests[counter_for_list_of_values_from_dictionary_automatic_tests]))
            run.bold = True
            run.italic = True

            # Add a new line to document
            prior_paragraph.add_run("\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run(
                'Number of negatives tests: ' + str(number_of_NOK_automatic_tests[counter_for_list_of_values_from_dictionary_automatic_tests]))
            run.bold = True
            run.italic = True

            # Add a new line to document
            prior_paragraph.add_run("\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run('Number of not run tests: '
                                          + str(len(list_of_automatic_tests[counter_for_list_of_values_from_dictionary_automatic_tests])
                                                - number_of_OK_automatic_tests[counter_for_list_of_values_from_dictionary_automatic_tests]
                                                - number_of_NOK_automatic_tests[counter_for_list_of_values_from_dictionary_automatic_tests]))
            run.bold = True
            run.italic = True

            # count all the white box tests
            counter_of_all_tests_automatic_tests += len(list_of_automatic_tests[counter_for_list_of_values_from_dictionary_automatic_tests])
            counter_for_list_of_values_from_dictionary_automatic_tests += 1

        # for Review coverage
        elif counter_for_list_of_values_from_dictionary_review_coverage < len(dict_with_all_data_from_review_coverage) and 'Tests Performed' in paragraph.text:
            # Add an empty paragraph
            prior_paragraph = insert_paragraph_after(paragraphs, idx + 1, "\r\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run('Number of tests: ' + str(
                len(dict_with_all_data_from_review_coverage[counter_for_list_of_values_from_dictionary_review_coverage][1])))
            run.bold = True
            run.italic = True

            # Add a new line to document
            prior_paragraph.add_run("\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run(
                'Number of positives tests: ' + str(number_of_OK_tests_review_coverage[counter_for_list_of_values_from_dictionary_review_coverage]))
            run.bold = True
            run.italic = True

            # Add a new line to document
            prior_paragraph.add_run("\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run(
                'Number of negatives tests: ' + str(number_of_NOK_tests_review_coverage[counter_for_list_of_values_from_dictionary_review_coverage]))
            run.bold = True
            run.italic = True

            # Add a new line to document
            prior_paragraph.add_run("\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run('Number of not run tests: '
                            + str(len(dict_with_all_data_from_review_coverage[counter_for_list_of_values_from_dictionary_review_coverage][1])
                                - number_of_OK_tests_review_coverage[counter_for_list_of_values_from_dictionary_review_coverage]
                                - number_of_NOK_tests_review_coverage[counter_for_list_of_values_from_dictionary_review_coverage]))
            run.bold = True
            run.italic = True

            # count all the review requirements tests
            counter_of_all_tests_review_coverage += len(dict_with_all_data_from_review_coverage[counter_for_list_of_values_from_dictionary_review_coverage][1])
            counter_for_list_of_values_from_dictionary_review_coverage += 1

        elif counter_for_list_of_values_from_dictionary_review_coverage >= len(dict_with_all_data_from_review_coverage):
            break

    # ADD TOTAL STATISTICS TO DOCUMENT

    global_summary_manual_tests_index = 0
    global_summary_automatic_tests_index = 0
    global_summary_rev_req_index = 0

    # Iterare and index all paragraphs
    for idx, paragraph in enumerate(paragraphs):
        # if found this paragraph keep the index
        if 'Global Summary for manual tests' in paragraph.text:
            global_summary_manual_tests_index = idx

        if 'Global Summary for automatic tests' in paragraph.text:
            global_summary_automatic_tests_index = idx

        if 'Global Summary for Reviews requirements tests' in paragraph.text:
            global_summary_rev_req_index = idx

        # for Manual tests
        if global_summary_manual_tests_index != 0 and idx < global_summary_manual_tests_index + 10:
            # if found a specific paragraph
            if 'Total number of tests:' in paragraph.text:
                # Add new text to it
                run = paragraph.add_run(str(counter_of_all_tests_manual_tests))
                run.bold = True
                run.italic = True

            if 'Total number of positives tests:' in paragraph.text:
                # Add new text to it
                run = paragraph.add_run(str(number_of_total_OK_manual_tests))
                run.bold = True
                run.italic = True

            if 'Total number of negatives tests:' in paragraph.text:
                # Add new text to it
                run = paragraph.add_run(str(number_of_total_NOK_manual_tests))
                run.bold = True
                run.italic = True

            if 'Total number of not run tests:' in paragraph.text:
                # Add new text to it
                run = paragraph.add_run(str(counter_of_all_tests_manual_tests - number_of_total_OK_manual_tests - number_of_total_NOK_manual_tests))
                run.bold = True
                run.italic = True

        # for Automatic tests
        if global_summary_automatic_tests_index != 0 and idx < global_summary_automatic_tests_index + 10:
            # if found a specific paragraph
            if 'Total number of tests:' in paragraph.text:
                # Add new text to it
                run = paragraph.add_run(str(counter_of_all_tests_automatic_tests))
                run.bold = True
                run.italic = True

            if 'Total number of positives tests:' in paragraph.text:
                # Add new text to it
                run = paragraph.add_run(str(number_of_total_OK_automatic_tests))
                run.bold = True
                run.italic = True

            if 'Total number of negatives tests:' in paragraph.text:
                # Add new text to it
                run = paragraph.add_run(str(number_of_total_NOK_automatic_tests))
                run.bold = True
                run.italic = True

            if 'Total number of not run tests:' in paragraph.text:
                # Add new text to it
                run = paragraph.add_run(
                    str(counter_of_all_tests_automatic_tests - number_of_total_OK_automatic_tests - number_of_total_NOK_automatic_tests))
                run.bold = True
                run.italic = True

        # forReview tests
        if global_summary_rev_req_index != 0 and idx < global_summary_rev_req_index + 10:
            # if found a specific paragraph
            if 'Total number of tests:' in paragraph.text:
                # Add new text to it
                run = paragraph.add_run(str(counter_of_all_tests_review_coverage))
                run.bold = True
                run.italic = True

            if 'Total number of positives tests:' in paragraph.text:
                # Add new text to it
                run = paragraph.add_run(str(number_of_total_OK_tests_review_coverage))
                run.bold = True
                run.italic = True

            if 'Total number of negatives tests:' in paragraph.text:
                # Add new text to it
                run = paragraph.add_run(str(number_of_total_NOK_tests_review_coverage))
                run.bold = True
                run.italic = True

            if 'Total number of not run tests:' in paragraph.text:
                # Add new text to it
                run = paragraph.add_run(
                    str(counter_of_all_tests_review_coverage - number_of_total_OK_tests_review_coverage - number_of_total_NOK_tests_review_coverage))
                run.bold = True
                run.italic = True

    document.save(path_to_generated_docx_file)
    try:
        update_toc(path_to_generated_docx_file)
    except:
        print("Table of content not updated!!!")


if __name__ == "__main__":
    main()