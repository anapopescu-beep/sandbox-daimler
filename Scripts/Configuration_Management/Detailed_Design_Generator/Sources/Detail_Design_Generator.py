"""Detail_Design_Generator.py: Used to generate detail design document.
Parameters has been given from argparse command line interface.
As input, it is needed component name, template of detail design
As output, it is needed a path to generate an docx file based upon given template
"""

__author__ = 'Ardeleanu Lucian'
__copyright__ = "Copyright 2021, Autoliv Electronic"
__version__ = "1.0"
__email__ = 'lucian.ardeleanu@autoliv.com'
__status__ = 'Released'

import os
import re
import argparse
from docx import Document
from docx.oxml.shared import OxmlElement, qn
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil
from subprocess import PIPE, Popen
import win32com.client as win32
from PIL import Image
import subprocess
import sys
from tqdm import tqdm
import yaml
from itertools import groupby
# ===============================================================
# FUNCTION USED TO REMOVE A ROW FROM TABLE
# INPUT ARGUMENT: OBJECT_TYPE TABLE AND ROW
# OUTPUT ARGUMENT: IT WILL REMOVE GIVEN ROW FROM TABLE
# ===============================================================
def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

# ===============================================================
# FUNCTION USED TO INSERT PARAGRAPH AFTER SPECIFIC PARAGRAPHS
# INPUT ARGUMENT: LIST OF OBJECT-BASED PARAGRAPHS, INTEGER: INDEX,
#                 STRING: TEXT TO INSERT
# OUTPUT ARGUMENT: IT WILL INSERT A PARAGRAPH INSIDE DOCUMENT
# ===============================================================
def insert_paragraph_after(paragraphs, idx, text=None):
    next_paragraph_idx = idx + 1
    if idx == len(paragraphs):
        return document.add_paragraph(text)
    next_paragraph = paragraphs[next_paragraph_idx]
    return next_paragraph.insert_paragraph_before(text)

# ===============================================================
# FUNCTION USED TO CONVERT A PNG IMAGE TO A BITMAP IMAGE
# This function will also convert a rectangular image to a sqare one.
# INPUT ARGUMENT: PATH TO IMAGE TO CONVERT
# OUTPUT ARGUMENT: PATH TO CONVERTED IMAGE
# ===============================================================
def convert_png_files_to_bmp_files_with_centering(file_to_convert):

    # ======================================================
    # CONVERT IMAGE FROM PNG TO BMP
    # ======================================================
    # replace in path only extension of file
    bmp_file_path = file_to_convert.replace('.png', '.bmp')

    # opem image and save it with other extension
    Image.open(file_to_convert).save(bmp_file_path)

    # ======================================================
    # Reshaping rectangular image to square
    # ======================================================
    # Open old image
    old_image = Image.open(bmp_file_path)

    # get it's dimensions
    old_size = old_image.size

    # Get max dimension between width and height
    max_dimension = max(old_size)

    if max_dimension <= 1500:
        # create a new image based on same dimension as maximum dimension from old image
        # first argument is type of image, second is dimension, third is white color ( to have a white background )
        new_size_of_new_image = (max_dimension, max_dimension)
        new_im = Image.new("RGB", new_size_of_new_image, (255, 255, 255))

        # add in center old image to noew created image
        new_im.paste(old_image, ((new_size_of_new_image[0] - old_size[0]) // 2, (new_size_of_new_image[1] - old_size[1]) // 2))

        # save new created image
        new_im.save(bmp_file_path)

        return bmp_file_path
    else:
        # print('\n', 'File ' + file_to_convert + ' has a very high resolution and could not been converted to bitmap!', '\n' )
        return file_to_convert

# ===============================================================
# FUNCTION USED TO CONVERT A PNG IMAGE TO A BITMAP IMAGE
# INPUT ARGUMENT: PATH TO IMAGE TO CONVERT
# OUTPUT ARGUMENT: PATH TO CONVERTED IMAGE
# ===============================================================
def convert_png_files_to_bmp_files(file_to_convert):

    # ======================================================
    # CONVERT IMAGE FROM PNG TO BMP
    # ======================================================
    # replace in path only extension of file
    bmp_file_path = file_to_convert.replace('.png', '.bmp')

    # opem image and save it with other extension
    Image.open(file_to_convert).save(bmp_file_path)

    # ======================================================
    # Reshaping rectangular image to square
    # ======================================================
    # Open old image
    old_image = Image.open(bmp_file_path)

    # get it's dimensions
    old_size = old_image.size

    # Get max dimension between width and height
    max_dimension = max(old_size)

    if max_dimension <= 1500:
        return bmp_file_path
    else:
        # print('\n', 'File ' + file_to_convert + ' has a very high resolution and could not been converted to bitmap!', '\n' )
        return file_to_convert


# ===============================================================
# FUNCTION USED TO COUNT ALL ELEMENTS IN A LIST
# INPUT ARGUMENT: LIST OF LISTS
# OUTPUT ARGUMENT: INT: COUNT OF ALL ELEMENTS
# ===============================================================
def get_elements_of_nested_list(element):
    count = 0
    if isinstance(element, list):
        for each_element in element:
            count += get_elements_of_nested_list(each_element)
    else:
        count += 1
    return count

def get_nested_length(list):
    size = 0
    for item in list:
        if type(item) == list:
            size += get_nested_length(item)
        else:
            size += 1
    return size

# ===============================================================
# FUNCTION USED TO move a table after a specific paragraph
# INPUT ARGUMENT: LIST OF OBJECT-BASED PARAGRAPHS, INTEGER: INDEX,
#                 STRING: TEXT TO INSERT
# OUTPUT ARGUMENT: IT WILL INSERT A PARAGRAPH INSIDE DOCUMENT
# ===============================================================
def move_table_after(table, paragraph):

    # Define a new paragraph entity
    new_p = OxmlElement("w:p")

    # Add table and blank paragraph
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)
    p.addnext(new_p)

    # Add a blank text
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(" ")

# ===============================================================
# FUNCTION USED TO ADD A PARAGRAPH AFTER A SPECIFIC PARAGRAPH
# INPUT ARGUMENT: OBJECT TYPE CELLS, STRING WITH COLOR IN HEXA
# OUTPUT ARGUMENT: NONE
# ===============================================================
def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para

# ===============================================================
# FUNCTION USED TO SHADE CELLS IN SPECIFIC COLLOR
# INPUT ARGUMENT: OBJECT TYPE CELLS, STRING WITH COLOR IN HEXA
# OUTPUT ARGUMENT: NONE
# ===============================================================
def shade_cells(cells, shade):
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcVAlign = OxmlElement("w:shd")
        tcVAlign.set(qn("w:fill"), shade)
        tcPr.append(tcVAlign)

# ===============================================================
# FUNCTIONS USED TO ADD NEW ENTRY'S IN TABLES
# INPUT ARGUMENT: OBJECTS
# OUTPUT ARGUMENT: NONE
# ===============================================================
def new_entry_function(row, text, table_obj):
    a = table_obj.rows[row].cells
    a[0].paragraphs[0].add_run(text).bold = True
    third_row_merge = (table_obj.cell(row, 0).merge((table_obj.cell(row, 1)))).merge(table_obj.cell(row, 3))
    shade_cells([table_obj.cell(row, 0)], "#000080")

def new_entry(row, text, table_obj):
    a = table_obj.rows[row].cells
    a[0].paragraphs[0].add_run(text).bold = True
    third_row_merge = (table_obj.cell(row, 0).merge((table_obj.cell(row, 1)))).merge(table_obj.cell(row, 2))
    shade_cells([table_obj.cell(row, 0)], "#000080")

def new_entry_simple_table(row, text, table_obj):
    a = table_obj.rows[row].cells
    a[0].paragraphs[0].add_run(text).bold = True
    shade_cells([table_obj.cell(row, 0)], "#000080")


def write_cell(row, cell, nr_to_merge, text, table_obj):
    table_obj.cell(row,cell).text = text
    if nr_to_merge != 0:
        while cell != nr_to_merge:
            row_merged = table_obj.cell(row, cell).merge((table_obj.cell(row, cell+1)))
            cell = cell + 1

def sub_entry(row, no_cells, text_input, color, table_obj):
    i = 0
    while i < no_cells:
        table_obj.cell(row, i).text = text_input.split()[i]
        shade_cells([table_obj.cell(row, i)], color)
        i = i + 1

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

# =======================================================================
# FUNCTION NEEDED TO APPLY CONTROL CHARS TO STRING
# Description: Needed because output from mandoc has control chars in it
#              In order to parse output from mandoc, all control chars
#              must be applied in string ( here, only backspace )
# INPUT ARGUMENT: STRING: STRING TO APPLY CONTROL CHARACTER
# OUTPUT ARGUMENT: STRING: STRING WITH CONTROL CHARACTER ADDED
# =======================================================================
def apply_backspace(s):
    while True:
        # if you find a character followed by a backspace, remove both
        t = re.sub('.\b', '', s, count=1)
        if len(s) == len(t):
            # now remove any backspaces from beginning of string
            return re.sub('\b+', '', t)
        s = t

# =======================================================================
# FUNCTION USED PARSE BETWEEN 2 GIVEN STRINGS A STRING
# INPUT ARGUMENT: STRINGS: strings to be parsed
# OUTPUT ARGUMENT: STRING: parsed strings
# =======================================================================
def parse_string_between_strings(string_to_parse, first_string):

    # pass the string to another string used for processing
    splitted_list = string_to_parse

    # define the next line with no identation in front
    next_line = ''

    # for line  in list fron first string to the end
    for line in splitted_list[ splitted_list.find(first_string) + len(first_string) :].split('\n')  :

        # compute leading spaces in front of line
        leading_spaces = len(line) - len(line.lstrip())

        # if there are no leading spaces and there is words in line
        if leading_spaces == 0 and len(line) > 0:

            # Retain next line with no leading spaces
            next_line = line
            break

    return string_to_parse[string_to_parse.find(first_string) + len(first_string): string_to_parse.find(next_line)]

# =======================================================================
# FUNCTION USED TO MODIFY A STRING AND APPLY CORECT IDENTIATION TO IT
# INPUT ARGUMENT: STRING: strings to be parsed
# OUTPUT ARGUMENT: STRING: parsed strings
# =======================================================================
def modify_string_for_correct_identation(string_to_modify):

    # Split string after enter char
    string_to_modify = string_to_modify.split('\n')

    # Define markers used in string preprocess
    modified_string = ''

    # for line and it's index in string
    for line, index in zip(string_to_modify, range(0, len(string_to_modify) - 1)):

        # Get next element
        next_element = string_to_modify[index + 1]

        # Compute leading space for it
        leading_spaces = len(line) - len(line.lstrip())
        leading_spaces_next_element = len(next_element) - len(next_element.lstrip())

        if leading_spaces == leading_spaces_next_element:
            modified_string += line
        else:
            modified_string += line + '\n'

    # Define final output string
    post_modified_string = ''

    # Split string after enter char
    modified_string = modified_string.split('\n')

    # For every line in modified string var
    for line in modified_string:

        # Get leading spaces in front of every line
        leading_spaces = len(line) - len(line.lstrip())

        # Construct a string with that leading spaces
        string_to_search = leading_spaces * ' '

        if 'Definition at line' not in line:
            # Replace leading spaces from the middle of the string, keeping leading spaces in front of string
            post_modified_string += string_to_search + ' '.join(line.split()) + '\n'

    # Return string
    return post_modified_string

# ===============================================================
# CLASS USED TO PARSE TEXT INPUT BY IDENTATION INTO A LIST
# INPUT ARGUMENT: STRING TO BE PARSED
# OUTPUT ARGUMENT: LIST WITH PARSED DATA
# ===============================================================
def call_class_needed_for_parse_string(string_to_parse):

    class Node:
        def __init__(self, indented_line):
            self.children = []
            self.level = len(indented_line) - len(indented_line.lstrip())
            self.text = indented_line.strip()

        def add_children(self, nodes):
            childlevel = nodes[0].level
            while nodes:
                node = nodes.pop(0)
                if node.level == childlevel:  # add node as a child
                    self.children.append(node)
                elif node.level > childlevel:  # add nodes as grandchildren of the last child
                    nodes.insert(0, node)
                    self.children[-1].add_children(nodes)
                elif node.level <= self.level:  # this node is a sibling, no more children
                    nodes.insert(0, node)
                    return

        def as_dict(self):
            if len(self.children) > 1:
                return [self.text, [node.as_dict() for node in self.children]]
            elif len(self.children) == 1:
                return [self.text, self.children[0].as_dict()]
            else:
                return self.text

    root = Node('')
    root.add_children([Node(line) for line in string_to_parse.splitlines() if line.strip()])
    d = root.as_dict()
    return(d)

# ==============================================================================================================================
#                                                       SCRIPT WORKFLOW
# ==============================================================================================================================
# Define path to config file
path_to_YAML_file = 'Detail_Design_Generator_Config_File.yaml'

# Read yaml config file
try:
    with open(path_to_YAML_file, 'r') as stream:
        yaml_config_file = yaml.load(stream, Loader=yaml.FullLoader)
except Exception as exception:
    print('YAML Config file could not been loaded. Exception Occured:', exception)
    print('This API will exit now!')
    exit(1)

# Set document path
doc_path = yaml_config_file['Path to template file']
doc = Document(doc_path)

# Set output generated docx path
if yaml_config_file['Path to output file'] == None:
    new_doc_path = os.path.join(os.getcwd(), yaml_config_file['Component to be analysed'] + "_SW_DetailedDesign" + ".docx" )
else:
    new_doc_path = os.path.join(yaml_config_file['Path to output file'], yaml_config_file['Component to be analysed'] + "_SW_DetailedDesign" + ".docx")

print('------------------------------------------')
print('     DETAILED DESIGN TOOL GENERATOR       ')
print('PLEASE MAKE SURE YOU HAVE S: DRIVE MAPPED!')
print('------------------------------------------')

# ======== PATH TO USED TOOLS ===============================================
# MANDOC TOOL:
path_to_mandoc_tool = os.getcwd() + r'\Tools\mandoc\bin\mandoc.exe '

# MORITZ TOOL:
path_to_moritz_tool = os.getcwd() + r'\Tools\Moritz\ansi_c_create.bat '
path_to_moritz_folder_src = os.getcwd() + r'\Tools\Moritz\src/'
path_to_moritz_folder_html = os.getcwd() + r'\Tools\Moritz\html/'

# DOXYGEN TOOL
path_to_doxygen_tool = os.getcwd() + r'\Tools\doxygen\bin\doxygen.exe'
path_to_doxygen_config_file = os.getcwd() + r'\Tools\doxygen\DOXYGEN_CONFIG_FILE'
path_to_doxygen_output = os.getcwd() + r'\Tools\doxygen\OUTPUT'
path_to_doxygen_man3_folder = path_to_doxygen_output + r'\man\man3/'

# GENERAL PATHS USED IN SCRIPT
applicaton_path = yaml_config_file['Path to Application Folder']
component_directory = os.path.join(applicaton_path, yaml_config_file['Component to be analysed'])
implementation_directory = os.path.join(component_directory, "Implementation")
design_directory = os.path.join(component_directory, "Design")

temp_directory_path = os.path.join(os.getcwd(), 'Temp')

# ===========================================================================

# =================== TO DO BEFORE RUNNING SCRIPT ===========================
# Broke execution if path S:/Components/Application does not exists!
if not os.path.exists(applicaton_path):
    print('PATH:',applicaton_path,  'has not been found! Check your path before running this tool again!')
    exit(1)

# task kill word
print('Word task will be killed forcefully!')
os.system('Taskkill /IM WINWORD.EXE /F')
# ============================================================================

# list of parsed functions
list_of_functions = []

# list of entities in groups
group_statistics = {}

# Init dict stuct and unions
structure_statistics = {}

# Statistics with functions, vars and macros
module_statistics = {}

def main():

    # CREATE TEMPORARY DIRECTORY TO STORE SOURCE FILES
    # ==================================================
    # If directory already exist
    if os.path.exists(temp_directory_path):
        for root, dirs, files in os.walk( temp_directory_path):

            # Remove all of it's contents
            for f in files:
                os.unlink(os.path.join(root, f))
            for d in dirs:
                shutil.rmtree(os.path.join(root, d))
    else:
        # Otherwise, create it
        os.mkdir(temp_directory_path)
    # ========================================================

    # COPY SOURCE CODE FILES FROM COMPONENT IN TEMP DIRECTORY
    # ========================================================================

    # Define flag used to check if a c_template file has been found
    template_found_flag = False


    # For every file in implementation dir
    for root, dirs, files in os.walk(design_directory):
        for file in files:

            # if a c_template file has been found
            if file.endswith('c_template'):

                print('File:', file, 'has been found, it will be proceseed.' )

                # Change flag
                template_found_flag = True

                # copy only c_template files in temp folder
                source_file_path = os.path.join(root, file)
                destination_file_path = os.path.join(temp_directory_path, file)
                shutil.copy(source_file_path, destination_file_path)

                # Rename c_template file to c file
                os.rename(destination_file_path, destination_file_path.replace('c_template','c'))

            # if a c_template file has been found
            if file.endswith('h_template'):

                print('File:', file, 'has been found, it will be proceseed.' )

                # Change flag
                template_found_flag = True

                # copy only c_template files in temp folder
                source_file_path = os.path.join(root, file)
                destination_file_path = os.path.join(temp_directory_path, file)
                shutil.copy(source_file_path, destination_file_path)

                # Rename c_template file to c file
                os.rename(destination_file_path, destination_file_path.replace('h_template','h'))

    # if no c_template file has been found
    if template_found_flag == False:

        print('No c_template or h_template files has been found, c and h files will be processed!')

        # for every file in implementation directory
        for root, dirs, files in os.walk(implementation_directory):
            for file in files:

                # copy c and h files in temp directory
                if file.endswith('.c') or file.endswith('.h') :

                    source_file_path = os.path.join(root, file)
                    destination_file_path = os.path.join(temp_directory_path, file)
                    shutil.copy(source_file_path, destination_file_path)

    # =============================================================================

    # RUN DOXYGEN FOR CURRENT MODULE
    # ====================================================================
    # Clear entire output folder of doxygen
    for root, dirs, files in os.walk(path_to_doxygen_output):
        for f in files:
            os.unlink(os.path.join(root, f))
        for d in dirs:
            shutil.rmtree(os.path.join(root, d))

    # ----------------- MODIFY THE DOXYGEN CONFIG FILE ----------------
    # Open DOXYGEN_CONFIG_FILE to write component to analyze and read file
    f = open(path_to_doxygen_config_file, "r")
    output_config_doxygen = f.readlines()
    f.close()

    # Set in config file input path of component and output path
    new_config_output = []
    for line in output_config_doxygen:
        # Add input path
        if 'INPUT                  = ' in line:
            new_config_output.append('INPUT                  = ' + temp_directory_path + '\n')

        # Add Output path
        elif 'OUTPUT_DIRECTORY       = ' in line:
            new_config_output.append('OUTPUT_DIRECTORY       = ' + path_to_doxygen_output + '\n')

        # Add new Aliases
        elif 'ALIASES                += ' in line:
            for tag_key, tag_value in yaml_config_file['New Tags'].items():

                # Construct string per every alias
                string_to_append = 'ALIASES                += ' + tag_key + '="\par ' + tag_value +  '^^"'  + '\n'

                # append string in config file
                if string_to_append not in new_config_output:
                    new_config_output.append(string_to_append)

        # Append rest of lines
        else:
            new_config_output.append(line)

    # Open DOXYGEN_CONFIG_FILE to write modified values
    f = open(path_to_doxygen_config_file, "w")
    f.writelines(new_config_output)
    f.close()

    print('================================================')
    print('Call Doxygen.exe in order to generate man3 files')

    # # ----------------- RUN DOXYGEN  ----------------
    subprocess.run(path_to_doxygen_tool + ' ' + path_to_doxygen_config_file + ' > ' + path_to_doxygen_output.replace('OUTPUT', '') + '\log_doxygen.txt  2>&1', shell=True)

    print('Doxygen.exe has been successfully executed!')
    print('================================================')

    # ----------------- REMODIFY THE DOXYGEN CONFIG FILE ----------------
    # Open DOXYGEN_CONFIG_FILE to write component to analyze and read file
    f = open(path_to_doxygen_config_file, "r")
    output_config_doxygen = f.readlines()
    f.close()

    # Set in config file input path of component and output path
    new_config_output = []
    for line in output_config_doxygen:
        # Add input path
        if 'INPUT                  = ' in line:
            new_config_output.append('INPUT                  = ' + '\n')

        # Add Output path
        elif 'OUTPUT_DIRECTORY       = ' in line:
            new_config_output.append('OUTPUT_DIRECTORY       = ' + '\n')

        # Add new Aliases
        elif 'ALIASES                += ' in line:
            # Construct string per every alias
            string_to_append = 'ALIASES                += ' + '\n'

            # append string in config file
            if string_to_append not in new_config_output:
                new_config_output.append(string_to_append)

        # Append rest of lines
        else:
            new_config_output.append(line)

    # Open DOXYGEN_CONFIG_FILE to write modified values
    f = open(path_to_doxygen_config_file, "w")
    f.writelines(new_config_output)
    f.close()

    # ====================================================================

    # COPY SOURCE FILES TO MORTIZ FOLDER TO GENERATE GRAPHICS
    # =================================================================
    # Delete all existing files from Moritz src
    for file in os.listdir(path_to_moritz_folder_src):
        os.remove(os.path.join(path_to_moritz_folder_src, file))

    # Copy source file into mortiz folder src
    for file in os.listdir(temp_directory_path):

        # Set source and destination path
        source_path = os.path.join(temp_directory_path, file)
        destination_path = path_to_moritz_folder_src + file

        # Copy files with shutil
        shutil.copy(source_path, destination_path)
    # =================================================================

    # =================================================================
    # CREATE LISTS OF SOURCE CODE AND MAN3 GENERATED FILES
    # =================================================================
    # d) List all files present at this location and save it into a table
    files_to_be_compared = []
    for file in os.listdir(temp_directory_path):
        files_to_be_compared.append(file)

    # e) Add .3 extension
    files_to_be_compared = [x + ".3" for x in files_to_be_compared]

    # 2. Sort and save in file_list all files with .c.3 extension
    file_list = []
    for file in os.listdir(path_to_doxygen_man3_folder):
        if file.endswith(".c.3") or file.endswith(".h.3") or file.endswith(".3"):
            file_list.append(file)

    list_of_man_files = []
    # 4. Make this verification while the counter is less tha the length of the files_to_be_compared size
    counter = 0
    while counter < len(files_to_be_compared):
        # a) Store in c_file the first(second...) entry of file_list (array which contains man3 sorted files)
        for man_file in file_list:
            # b) Store in first the first element from files_to_be_compared table (src/ic files from application path)
            for first in files_to_be_compared:
                # c) If the both file names are equal
                if man_file == first or ('Group' in man_file and yaml_config_file['Component to be analysed'] in man_file):
                    # Append in a list all c files found
                    list_of_man_files.append(man_file)
                    counter = counter + 1

    # ===============================================================
    # PROCESS MAN3 FILES WITH STRUCT OR UNION IT IT
    # ===============================================================
    for file in file_list:
        if (file not in list_of_man_files) and ('Union' in file or 'Struct' in file ):

            structure_statistics[file] = {}

            # Define list with data per structure
            structure_data = []

            # ============= CALL MANDOC TOOL TO OUTPUT MAN3 OUTPUT FILE ===================
            # Execute command to extract data from man file using mandoc tool
            command = path_to_mandoc_tool + path_to_doxygen_man3_folder + file

            # Call command using subprocess and filter results
            with Popen(command, stdout=PIPE, stderr=None, shell=True) as process:
                output = process.communicate()[0].decode("utf-8").split('\n')

            # ========= APPLY CONTROL CHARS TO OUTPUTED STRING ========================
            # Define a new list with strings to append data
            output_list_of_strings = []
            string_with_output_unprocessed = " "

            for elem in output:
                # Convert and append data in new list
                output_list_of_strings.append(str(apply_backspace(elem)))

                # Append it into a list, in case it's needed
                string_with_output_unprocessed += str(apply_backspace(elem))

            # process string a little bit by changing line terminators
            string_from_output = string_with_output_unprocessed.replace('\r', '\n')

            if 'Detailed Description' in string_from_output:
                description_part = parse_string_between_strings(string_to_parse=string_from_output,first_string='Detailed Description')
                description_modified_string = modify_string_for_correct_identation(description_part)
            else:
                description_modified_string = '-'

            if 'Field Documentation' in string_from_output:
                field_documentation_part = parse_string_between_strings(string_to_parse=string_from_output,first_string='Field Documentation')
                field_documentation_modified_string = modify_string_for_correct_identation(field_documentation_part)
            else:
                field_documentation_modified_string = '-'

            try:
                # call class to extract string with description
                description_string = call_class_needed_for_parse_string(description_modified_string)[1]

                # call class to extract string with variables and them description
                field_documentation_list = call_class_needed_for_parse_string(field_documentation_modified_string)[1]
            except Exception as e:
                print('Struct or union', file, 'could not been parsed. Maybe you not wrote it accordingly. Exception Occured:', e)
                exit(1)

            # Init a global variable list
            global_field_documentation_list = []

            # If found more than one variable and it's in a list
            if len(field_documentation_list) > 1:
                if isinstance(field_documentation_list[0], list):
                    for elem in field_documentation_list:
                        global_field_documentation_list.append(elem)
                if isinstance(field_documentation_list[0], str) and len(field_documentation_list) > 2:
                    for elem in field_documentation_list:
                        global_field_documentation_list.append(elem)
                if isinstance(field_documentation_list[0], str) and len(field_documentation_list) == 2:
                    global_field_documentation_list.append(field_documentation_list)
            else:
                global_field_documentation_list.append(field_documentation_list)

            # assign in dictionary description
            if len(description_string) != 0:
                structure_statistics[file] = {
                    "Description" : [description_string],
                    "Fields": global_field_documentation_list,
                }

    # ===========================================================================
    # PROCESS MAN3 FILES WITH SOURCE CODE
    # ===========================================================================
    for man_file in list_of_man_files:
        module_statistics[man_file] = {
            'macros': [],
            'functions': [],
            'variables': [],
            'includes': [],
        }

        # ============= CALL TOOL TO OUTPUT MAN3 OUTPUT FILE ===================
        # Execute command to extract data from man file using mandoc tool
        command = path_to_mandoc_tool + path_to_doxygen_man3_folder + man_file

        # Call command using subprocess and filter results
        with Popen(command, stdout=PIPE, stderr=None, shell=True) as process:
            output = process.communicate()[0].decode("utf-8").split('\n')

        # ========= APPLY CONTROL CHARS TO OUTPUTED STRING ========================
        # Define a new list with strings to append data
        output_list_of_strings = []
        string_with_output_unprocessed = " "

        for elem in output:
            # Convert and append data in new list
            output_list_of_strings.append(str(apply_backspace(elem)))

            # Append it into a list, in case it's needed
            string_with_output_unprocessed += str(apply_backspace(elem))

        # process string a little bit by changing line terminators
        string_from_output = string_with_output_unprocessed.replace('\r', '\n')

        # ONLY FOR DEBUG FUNCTION
        # f = open(man_file + "_raw.txt", "w")
        # f.writelines(string_from_output)
        # f.close()
        # ONLY FOR DEBUG FUNCTION

        # ======================= PARSE TEXT FILES IN SPECIFIC STRINGS ==================================
        # extract includes from entire string
        if 'Function Documentation' in string_from_output:
            functions_part = parse_string_between_strings(string_to_parse=string_from_output, first_string='Function Documentation')
            function_modified_string = modify_string_for_correct_identation(functions_part)
        else:
            function_modified_string = '-'

        if 'Macro Definition Documentation' in string_from_output:
            defines_part = parse_string_between_strings(string_to_parse=string_from_output,first_string='Macro Definition Documentation')
            defines_modified_string = modify_string_for_correct_identation(defines_part)
        else:
            defines_modified_string = '-'

        # extract includes from entire string
        if 'Variable Documentation' in string_from_output:
            variable_part = parse_string_between_strings(string_to_parse=string_from_output, first_string='Variable Documentation')
            variable_modified_string = modify_string_for_correct_identation(variable_part)
        else:
            variable_modified_string = '-'


        # ONLY FOR DEBUG FUNCTION
        # f = open(man_file + "log_output_conversion.txt", "w")
        # f.writelines(variable_modified_string)
        # f.close()
        # ONLY FOR DEBUG FUNCTION

        # =============================================================================
        # GROUP VARIABLES INTO SINGLE LIST
        # =============================================================================
        # Call class for variable list
        variables_list = call_class_needed_for_parse_string(variable_modified_string)[1]

        # Init a global variable list
        global_variable_list = []

        # If found more than one variable and it's in a list
        if len(variables_list) > 1:
            if isinstance(variables_list[0], list):
                for var in variables_list:
                    global_variable_list.append(var)
            if isinstance(variables_list[0], str) and len(variables_list) > 2:
                for var in variables_list:
                    global_variable_list.append(var)
            if isinstance(variables_list[0], str) and len(variables_list) == 2:
                global_variable_list.append(variables_list)
        else:
            global_variable_list.append(variables_list)

        # =========================================================
        # GROUP FUNCTIONS INTO SINGLE LIST
        # =========================================================

        # Call class for variable list
        function_list = call_class_needed_for_parse_string(function_modified_string)[1]
        # Init a global variable list
        global_function_list = []

        # If found more than one variable and it's in a list
        if len(function_list) > 1:
            if isinstance(function_list[0], list):
                for func in function_list:
                    global_function_list.append(func)
            if isinstance(function_list[0], str) and isinstance(function_list[1], str) :
                for func in function_list:
                    global_function_list.append(func)
            if isinstance(function_list[0], str) and isinstance(function_list[1], list):

                global_function_list.append(function_list)
        else:
            global_function_list.append(function_list)


        # =========================================================
        # GROUP DEFINES INTO SINGLE LIST
        # =========================================================
        # Call class for variable list
        macro_list = call_class_needed_for_parse_string(defines_modified_string)[1]

        # Init a global variable list
        global_macro_list = []

        # If found more than one variable and it's in a list
        if len(macro_list) > 0:

            # case of macro and description
            if isinstance(macro_list[0], list):
                for macro in macro_list:
                    global_macro_list.append(macro)

            if isinstance(macro_list[0], str):
                # case of a single macro and single description
                if '#define' in macro_list[0] and len(macro_list) == 2 and '#define' not in macro_list[1] :
                    global_macro_list.append(macro_list)

                # multiple macros without description
                if len(macro_list) >= 2:
                    for macro in macro_list:
                        global_macro_list.append(macro)


        # ASSERT LISTS IN DICTIONARY
        module_statistics[man_file] = {
            'macros': global_macro_list,
            'functions': global_function_list,
            'variables': global_variable_list,
        }

        # ======================================================================
        # GROUP DEFINING AND APPENDING FILE AND GROUP DESCRIPTION IN GROUP DICT
        # ======================================================================
        if ('Group' in man_file) and (yaml_config_file['Component to be analysed'] in man_file):

            # Init variable for group description
            group_description_string = ''

            # =============== PARSE OUTPUT OF MANDOC TOOL FROM MAN3 FILE =====================

            # SET ONLY PIECE OF STRING YOU NEED TO PARSE FROM ENTIRE OUTPUT OF MANDOC
            if 'Detailed Description' in string_from_output:

                # parse group description
                group_description = parse_string_between_strings(string_to_parse=string_from_output,first_string='Detailed Description')

                # MODIFY STRING FOR CORRENT IDENTATION
                group_description = modify_string_for_correct_identation(group_description)

                # CALL CLASS TO OUTPUT MODIFIED STRING
                group_description_string = call_class_needed_for_parse_string(group_description)[1]

            # ========= ADD DESCRIPTION IN GROUP STATISTICS DICTIONARY =======================
            # If current file is already in dictionary
            if man_file in group_statistics:
                # If string from output of mandoc has a description
                if len(group_description_string) > 0:
                    # Add description in dict
                    group_statistics[man_file]['Description'] = [group_description_string]
                else:
                    group_statistics[man_file]['Description'] = ['None']

            # Otherwise, create a new entry for file
            else:
                group_statistics[man_file] = {}

                # Add description in dict
                if len(group_description_string) > 0:
                    group_statistics[man_file]['Description'] = [group_description_string]
                else:
                    group_statistics[man_file]['Description'] = ['None']

    # Remove temp directory
    shutil.rmtree(temp_directory_path)

    # For debug purpose
    # for man_file in module_statistics:
    #     for key,value in module_statistics[man_file].items():
    #         print(key,value)

    # for man_file in group_statistics:
    #     print(man_file)
    #     for key,value in group_statistics[man_file].items():
    #         print(key,value)

def function_parser():
    print('PARSING FUNCTIONS')

    print('=====================================================')
    print('Call Moritz batch files in order to generate graphics')
    subprocess.run(path_to_moritz_tool + '> ' + path_to_moritz_folder_src + 'log_moritz.txt  2>&1', shell=False)
    print('Moritz graphs has been successfully generated!')
    print('=====================================================')


    doc.add_heading("Services", level=2)
    set_nr_rows = 20
    parsed_function_name = ''
    sort_list = []

    for file in module_statistics:

        # Define list where group functions will be appended
        group_functions_list_per_file = []

        for x in module_statistics[file]:
            if x == 'functions':
                # first name is function name, second name is function description, if exists
                # for test_case_with_path, list_increment in zip( value, tqdm(value) ):
                print('\n', "Parsing File:", file )
                for function, list_increment in zip( module_statistics[file][x], tqdm(module_statistics[file][x]) ):

                    # MARKER INDICATORS FOR FILLING TABLES
                    # ==================================================================
                    function_definition = ''
                    function_description = 'The function is called by *** to ***'
                    function_description_parameters = 'NA'
                    function_description_returns = 'NA'
                    function_dynamic_aspect_caller = '*'
                    function_dynamic_aspect_description = '*'
                    function_static_aspect = '*'
                    function_constrains = ' '
                    function_input_parameters = 'Name: NA; Type: NA; Description: NA; Range: NA;'
                    function_output_parameters = 'Name: NA; Type: NA; Description: NA; Range: NA;'


                    list_of_parameters_from_a_specific_position = ['InputParameters', 'OutputParameters', 'Returns', 'DynamicAspectDescription', 'DynamicAspectCaller', 'StaticAspect', 'Constrains' ]

                    lines_to_add_in_table = []


                    if len(function) != 0 and function != '-':
                        # [ HERE WE ARE, [ .... ] ]
                        # Check if first element is a string
                        if isinstance(function[0], str) and len(function[0]) > 5:
                            # if it is a string, then it's a function definition
                            function_definition = function[0]

                        # If there are multiple elements than a function definition
                        if len(function) > 1:

                            # check first element
                            if isinstance(function[1], list ):
                                for element in function[1]:

                                    # [ function_definition, [ HERE WE ARE [...] ] ]
                                    # and first element from list is a string with multiple words
                                    if isinstance(element, str) and len(element.split(' ')) > 1:
                                        function_description = element

                                    if isinstance(element, list) and len(element) == 2 and isinstance(element[0], str) and isinstance(element[1], str) :

                                        # PARAMETER NOT USED
                                        if  ( len(element[0].split(' ')) == 1 and 'Parameter' in element[0] ) or ( len(element[1].split(' ')) == 1 and 'Parameter' in element[1] ):
                                            function_description_parameters = element[1]

                                        if ( len(element[0].split(' ')) == 1 and 'Return' in element[0] ) or ( len(element[1].split(' ')) == 1 and 'Return' in element[1] ):
                                            function_description_returns = element[1]

                                        if  ( len(element[0].split(' ')) == 1 and 'DynamicAspectDescription' in element[0] ) or ( len(element[1].split(' ')) == 1 and 'DynamicAspectDescription' in element[1] ):
                                            function_dynamic_aspect_description = element[1]

                                        if  ( len(element[0].split(' ')) == 1 and 'DynamicAspectCaller' in element[0] ) or ( len(element[1].split(' ')) == 1 and 'DynamicAspectCaller' in element[1] ):
                                            function_dynamic_aspect_caller = element[1]

                                        if  ( len(element[0].split(' ')) == 1 and 'StaticAspect' in element[0] ) or ( len(element[1].split(' ')) == 1 and 'StaticAspect' in element[1] ):
                                            function_static_aspect = element[1]

                                        if  ( len(element[0].split(' ')) == 1 and 'Constrains' in element[0] ) or ( len(element[1].split(' ')) == 1 and 'Constrains' in element[1] ):
                                            function_constrains = element[1]

                                        if  ( len(element[0].split(' ')) == 1 and 'InputParameters' in element[0] ) or ( len(element[1].split(' ')) == 1 and 'InputParameters' in element[1] ):
                                            function_input_parameters = element[1]

                                        if  ( len(element[0].split(' ')) == 1 and 'OutputParameters' in element[0] ) or ( len(element[1].split(' ')) == 1 and 'OutputParameters' in element[1] ):
                                            function_output_parameters = element[1]

                                        if ( ( len(element[0].split(' ')) == 1 ) or ( len(element[1].split(' ')) == 1 ) ) and (element[0] not in list_of_parameters_from_a_specific_position ):
                                            lines_to_add_in_table.append(element)

                    # ============================================================
                    # PARSE FUNCTION TO EXTRACT IT'S NAME
                    # ============================================================
                    # string to retain the function name
                    function_name = ''
                    if function_definition != '':

                        # parse function name using regex
                        regex1 = re.search('(.*)(.*?\s+)(\()(.*)(\))(.*)', function_definition)
                        if regex1 != None:
                            a = (len(regex1.group(1).split())) - 1
                            parsed_function_name = (regex1.group(1).split()[int(a)])

                    # sort list of parsed functions
                    if ( parsed_function_name not in sort_list) :
                        sort_list.append(parsed_function_name)
                        FUNC_search = re.search('(^FUNC)', function_definition)
                        PVAR = re.search('P2VAR', function_definition)
                        if (FUNC_search != None):
                            print()
                        else:
                            list_elem_output = []
                            list_elem_input = []

                            #print(z)
                            # Add Header
                            regey = re.search('ISR', function_definition)

                            # EXTRACT FUNCTION NAME FROM DICTIONARY
                            if regey != None:
                                regex11 = re.search('(\()(.*)(\))(.*)', function_definition)
                                if regex11 != None and '\&' not in function_definition:
                                    doc.add_heading(regex11.group(2), level=3)
                                    function_name = regex11.group(2)
                            else:
                                regex1 = re.search('(.*)(.*?\s+)(\()(.*)(\))(.*)', function_definition)
                                if regex1 != None and '\&' not in function_definition :
                                    a = (len(regex1.group(1).split())) - 1
                                    doc.add_heading(regex1.group(1).split()[int(a)], level=3)
                                    function_name = regex1.group(1).split()[int(a)]

                    # ============================================================

                        # if there is a string in function name
                        if function_name != '':

                            # b) Create a table and get the alv style and init raw counter
                            table = doc.add_table(rows=0, cols=4)
                            table.style = 'Table Grid'
                            curr_raw = 0
                            # ===========================================================

                            # WRITE DESCRIPTION IN DOCUMENT
                            # ===========================================================
                            table.add_row()
                            new_entry_function(curr_raw, "Object", table)
                            curr_raw += 1

                            # Write function description
                            table.add_row()
                            write_cell(curr_raw, 0, 3, function_description, table)
                            curr_raw += 1
                            # ============================================================

                            # WRITE PROTOTYPE IN DOCUMENT
                            # ===========================================================
                            # c) Write the function prototype
                            table.add_row()
                            new_entry_function(curr_raw, "Prototype", table)
                            curr_raw += 1

                            table.add_row()
                            write_cell(curr_raw, 0, 3, function_definition, table)
                            curr_raw += 1
                            # ===========================================================

                            # WRITE DIFFERENT PROPERTIES FOUND PARSING MAN3 FILES
                            # ===========================================================
                            # If properties found
                            if len(lines_to_add_in_table) > 0:

                                # For every property in appended list
                                for element in lines_to_add_in_table:

                                    # Write property name
                                    table.add_row()
                                    new_entry_function(curr_raw, element[0], table)
                                    curr_raw += 1

                                    # Write property value
                                    table.add_row()
                                    write_cell(curr_raw, 0, 3, element[1], table)
                                    curr_raw += 1

                            # ===========================================================

                            # WRITE INPUT PARAMETERS OF FUNCTION IN TABLE
                            # ====================================================================
                            # d) Write the input description
                            table.add_row()
                            new_entry_function(curr_raw, "Input parameters", table)
                            curr_raw += 1

                            table.add_row()
                            sub_entry(curr_raw, 4, "Name Type Description Range", "#C6D9F1", table)
                            curr_raw += 1

                            # Split after finding Name tag in multiple lists of parameters
                            splitted_parameters = function_input_parameters.split('Name: ')

                            # Filter empty elements from list
                            splitted_parameters = list(filter(None, splitted_parameters))

                            # Add back Name parameter ( removed after upper split operation )
                            splitted_parameters = ['Name: ' + string for string in splitted_parameters]

                            for parameter in splitted_parameters:
                                table.add_row()
                                splitted_parameter = parameter.split(';')
                                for parameter in splitted_parameter:
                                    if 'Name:' in parameter:
                                        table.cell(curr_raw, 0).text = parameter[parameter.find('Name: ') + len('Name: ') :]
                                    if 'Type:' in parameter:
                                        table.cell(curr_raw, 1).text = parameter[parameter.find('Type: ') + len('Type: ') :]
                                    if 'Description:' in parameter:
                                        table.cell(curr_raw, 2).text = parameter[parameter.find('Description: ') + len('Description: '):]
                                    if 'Range:' in parameter:
                                        table.cell(curr_raw, 3).text = parameter[parameter.find('Range: ') + len('Range: '):]
                                curr_raw = curr_raw + 1

                            # WRITE OUTPUT PARAMETERS OF FUNCTION IN TABLE
                            # ====================================================================

                            # f) Write the output description
                            table.add_row()
                            new_entry_function(curr_raw, "Output parameters", table)
                            curr_raw = curr_raw + 1

                            table.add_row()
                            sub_entry(curr_raw, 4, "Name Type Description Range", "#C6D9F1", table)
                            # If list_elem_output != (output parameters exist) write them in docx file
                            curr_raw = curr_raw + 1

                            # Split after finding Name tag in multiple lists of parameters
                            splitted_parameters = function_output_parameters.split('Name: ')

                            # Filter empty elements from list
                            splitted_parameters = list(filter(None, splitted_parameters))

                            # Add back Name parameter ( removed after upper split operation )
                            splitted_parameters = ['Name: ' + string for string in splitted_parameters]

                            for parameter in splitted_parameters:
                                table.add_row()
                                splitted_parameter = parameter.split(';')
                                for parameter in splitted_parameter:
                                    if 'Name:' in parameter:
                                        table.cell(curr_raw, 0).text = parameter[parameter.find('Name: ') + len('Name: ') :]
                                    if 'Type:' in parameter:
                                        table.cell(curr_raw, 1).text = parameter[parameter.find('Type: ') + len('Type: ') :]
                                    if 'Description:' in parameter:
                                        table.cell(curr_raw, 2).text = parameter[parameter.find('Description: ') + len('Description: '):]
                                    if 'Range:' in parameter:
                                        table.cell(curr_raw, 3).text = parameter[parameter.find('Range: ') + len('Range: '):]
                                curr_raw = curr_raw + 1


                            # WRITE RETURN IN TABLE
                            # =================================================================================
                            table.add_row()
                            new_entry_function(curr_raw, "Return value", table)
                            curr_raw += 1

                            table.add_row()
                            sub_entry(curr_raw, 2, "Type Description", "#C6D9F1", table)

                            a = table.cell(curr_raw, 1).merge(table.cell(curr_raw, 3))
                            curr_raw += 1

                            # Write the actual return value of the function
                            regex11 = re.search('(.*)(\()', function[0])

                            if regex11 != None:
                                a = (len(regex11.group(1).split())) - 2
                                if regex11.group(1).split()[int(a)] != "void":
                                    table.add_row()
                                    write_cell(curr_raw, 0, 0, regex11.group(1).split()[int(a)], table)
                                    write_cell(curr_raw, 1, 3, function_description_returns, table)
                                else:
                                    table.add_row()
                                    write_cell(curr_raw, 0, 0, "NA", table)
                                    write_cell(curr_raw, 1, 3, function_description_returns, table)
                            curr_raw += 1
                            # ================================================================================
                            table.add_row()
                            new_entry_function(curr_raw, "Dynamic aspect", table)
                            curr_raw += 1

                            table.add_row()
                            sub_entry(curr_raw, 2, "Who(callers) Description", "#C6D9F1", table)
                            b = table.cell(curr_raw, 1).merge(table.cell(curr_raw, 3))
                            curr_raw += 1

                            table.add_row()
                            write_cell(curr_raw, 0, 0, function_dynamic_aspect_caller, table)
                            write_cell(curr_raw, 1, 0, function_dynamic_aspect_description, table)
                            b = table.cell(curr_raw, 1).merge(table.cell(curr_raw, 3))
                            curr_raw += 1

                            table.add_row()
                            new_entry_function(curr_raw, "Static aspect", table)
                            curr_raw += 1

                            table.add_row()
                            write_cell(curr_raw, 0, 3, function_static_aspect, table)
                            curr_raw += 1

                            table.add_row()
                            new_entry_function(curr_raw, "Constrains", table)
                            curr_raw += 1

                            table.add_row()
                            write_cell(curr_raw, 0, 3, function_constrains, table)

                            # append in global function list function name
                            list_of_functions.append(function_name)

                            if 'Group' in file:
                                group_functions_list_per_file.append(function_name)

                            # ==========================================================================
                            # ADD PARAGRAPH FOR IDENTIFY WHERE TO INSERT GRAPHICS FROM MORITZ
                            # ==========================================================================
                            list_with_all_files_in_directory = os.listdir(path_to_moritz_folder_html)

                            paragraph_to_insert = ''

                            # list all files in html dir
                            for moritz_file in list_with_all_files_in_directory:

                                # if Moritz generated our function graphic
                                if 'dot_' in moritz_file and function_name in moritz_file and moritz_file.endswith('.png'):

                                    # Construct a string used by another function to figure out where to insert graphs
                                    paragraph_to_insert = 'dot_' + function_name + '.png'

                            # add string in document
                            if paragraph_to_insert != '':
                                # if exists, add it's name in Word, in order to know where to insert graphics from Mortiz tool
                                paragraph = doc.add_paragraph(paragraph_to_insert)
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                paragraph = doc.add_paragraph('')
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # APPEND FUNCTION NAME INTO SPECIAL DICT IF FUNCTION HAS BEEN FOUND IN A GROUP
        # ============================================================================
        if group_functions_list_per_file != []:
            if file in group_statistics:
                group_statistics[file]['Functions'] = group_functions_list_per_file
            else:
                group_statistics[file] = {}
                group_statistics[file]['Functions'] = group_functions_list_per_file


    doc.save(new_doc_path)

def variable_parser():

    print('PARSING VARIABLES')

    sort_list = []
    doc.add_heading("Variabiles", level=2)

    variable_name = ''

    for file in module_statistics:

        # Define a list for variables group
        group_variables_list_per_file = []

        for x in module_statistics[file]:
            if x == 'variables':
                print('\n', "Parsing File:", file )
                # for test_case_with_path, list_increment in zip( value, tqdm(value) ):
                for variable, list_increment in zip( module_statistics[file][x], tqdm(module_statistics[file][x]) ):

                    variable_definition = ''
                    variable_description = "This variable is used to ..."
                    lines_to_add_in_table = []


                    if len(variable) != 0 and variable != '-':
                        # [ HERE WE ARE, [ .... ] ]
                        # Check if first element is a string

                        # if there is a single variable
                        if isinstance(variable, str) and len(variable) > 5:
                            variable_definition = variable

                        # If there is a group of variable and description
                        if isinstance(variable, list) and all(isinstance(item, str) for item in variable) and len(variable) == 2:
                            variable_definition = variable[0]
                            variable_description = variable[1]

                        # If there is a group with multiple elements than definition and description
                        # check first element
                        if isinstance(variable, list ) and isinstance(variable[0], str) and isinstance(variable[1], list):

                            # Define variable name
                            variable_definition = variable[0]

                            for element in variable[1]:
                                # [ function_definition, [ HERE WE ARE [...] ] ]
                                # and first element from list is a string with multiple words
                                if isinstance(element, str) and len(element.split(' ')) > 1:
                                    variable_description = element

                                if isinstance(element, list) and len(element) == 2 and isinstance(element[0], str) and isinstance(element[1], str) :
                                    lines_to_add_in_table.append(element)

                    # ================================================================================
                    # PARSING VARIABLE DEFINITION FOR EXTRACT HEADING TITLE
                    # ================================================================================
                    if variable_definition not in sort_list and variable_definition != '' :
                        if(len(variable_definition.split())) > 3:

                            regex = re.search('=', variable_definition)
                            if regex != None:
                                regex2 = re.search('(. )(.*)(=)', variable_definition)
                                if len(regex2.group(2).split()) > 2:
                                    regex3 = re.search('(.*)(\[)(.*)(\])', regex2.group(2))
                                    if regex3 != None:
                                        # verifica daca e array
                                        regex4 = re.search('(.*)(.*?\s+(\S+))', regex3.group(1))
                                        if regex4 != None:
                                            doc.add_heading(regex4.group(2), level=3)
                                            variable_name = regex4.group(2)

                                    else:
                                        doc.add_heading(regex2.group(2).split()[2], level=3)
                                        variable_name = regex2.group(2).split()[2]
                                else:
                                    if len(regex2.group(2).split()) > 1:
                                        doc.add_heading(regex2.group(2).split()[1], level=3)
                                        variable_name = regex2.group(2).split()[1]
                                    else:
                                        doc.add_heading(regex2.group(2), level=3)
                                        variable_name = regex2.group(2)
                            else:
                                regex5 = re.search('(.*?\])', variable_definition)
                                if regex5 != None:
                                    regex3 = re.search('(.*)(.*?\s+(\S+))', variable_definition)
                                    if regex3 != None:
                                        if '[' not in regex5.group(1) or ']' not in regex5.group(1):
                                            regex4 = re.search('(.*)(.*?\s+(\S+))', regex3.group(1))
                                            if regex4 != None:
                                                doc.add_heading(regex4.group(2), level=3)
                                                variable_name = regex4.group(2)
                                        else:
                                            if 'extern' in regex3.group(2):
                                                continue
                                            else:
                                                regex_with_parathesis = re.search('.*\s(.*?\[.*\])', variable_definition)
                                                variable_name = regex_with_parathesis.group(1)
                                                doc.add_heading(variable_name, level=3)
                                else:
                                    regex6 = re.search('(.*)(.*?\s+(\S+))', variable_definition)
                                    if regex6 != None:
                                        doc.add_heading(regex6.group(2), level=3)
                                        variable_name = regex6.group(2)

                        elif (len(variable_definition.split())) == 2:
                            doc.add_heading(variable_definition.split()[1], level=3)
                            variable_name = variable_definition.split()[1]

                        elif (len(variable_definition.split())) == 3:

                            if 'LOCAL' in variable_definition:
                                # extract variable name
                                extracted_variable_name = variable_definition.split()[-1]
                            else:
                                # extract variable name
                                extracted_variable_name = variable_definition.split()[-2]

                            # insert it into document
                            doc.add_heading(extracted_variable_name, level=3)
                            variable_name = extracted_variable_name
                            # regex5 = re.search('(.*?\])', z)
                            # if regex5 != None:
                            #     regex3 = re.search('(.*)(.*?\s+(\S+))', z)
                            #     if regex3 != None:
                            #         regex4 = re.search('(.*)(.*?\s+(\S+))', regex3.group(1))
                            #         if regex4 != None:
                            #             doc.add_heading(regex4.group(2), level=3)
                        else:
                            doc.add_heading(variable_definition, level=3)
                            variable_name = variable_definition

                        # ================================================================================

                        if variable_name != '':

                            # ADD TABLE AND ROWS
                            # ========================================================
                            # Define nr of rows
                            nr_rows = 6 + get_elements_of_nested_list(lines_to_add_in_table)

                            # Add table in document
                            table = doc.add_table(rows=nr_rows, cols=3)

                            # Set table style
                            table.style = 'Table Grid'
                            # ========================================================

                            # ========================================================
                            # ADD VARIABLE TYPE
                            # ========================================================
                            write_cell(0, 0, 0, "Type", table)
                            #print(z)
                            if (len(variable_definition.split())) > 3:
                                regex = re.search('=', variable_definition)

                                if regex != None:#daca am egal
                                    regex2 = re.search('(^[A-Z]\w+)(.*)(=)(.*)', variable_definition)
                                    #daca am litera mare si egal
                                    if regex2 != None:
                                        regex3 = re.search('(.*)(\[)(.*)(\])', regex2.group(2))
                                        if regex3 != None:
                                            #verifica daca e array
                                            regex4 = re.search('(.*)(.*?\s+(\S+))', regex3.group(1))
                                            if regex4 != None:
                                                write_cell(1, 0, 0, regex4.group(1), table)
                                            #nu e array
                                        else:
                                            regex5 = re.search('(.*)(.*?\s+(\S+))', regex2.group(2))
                                            if regex5 != None:
                                                write_cell(1, 0, 0, regex5.group(1), table)
                                    #daca am litera mica si egal
                                    regex30 = re.search('((^[a-z]\w+)(.*))(=)(.*)', variable_definition)
                                    if regex30 != None:
                                        #verifica daca e array
                                        regex3 = re.search('(.*)(\[)(.*)(\])', regex30.group(1))
                                        if regex3 != None:
                                            regex4 = re.search('(.*)(.*?\s+(\S+))', regex3.group(1))
                                            if regex4 != None:
                                                write_cell(1, 0, 0, regex4.group(1), table)
                                        else:
                                            regex6 = re.search('(.*)(.*?\s+(\S+))', regex30.group(1))
                                            if regex6 != None:
                                                write_cell(1, 0, 0, regex6.group(1), table)

                                else:#daca nu am egal si am litera mare
                                    regex7 = re.search('(^[A-Z]\w+)(.*)(.*?\s+(\S+))', variable_definition)
                                    if regex7 != None:
                                        #daca am array
                                        regex301 = re.search('(.*)(\[)(.*)(\])', variable_definition)
                                        if regex301 != None:
                                            regex40 = re.search('(^[A-Z]\w+)(.*)(.*?\s+(\S+))', regex301.group(1))
                                            if regex40 != None:
                                                # print(regex40.group(2))
                                                write_cell(1, 0, 0, regex40.group(2), table)
                                        #regex9 = re.search('(.*)(.*?\s+(\S+))', regex7.group(2))
                                        #if regex9 != None:
                                            #write_cell(1, 0, 0, regex9.group(1), table)
                                        else:
                                            write_cell(1, 0, 0, regex7.group(2), table)

                                #nu am egal si am litera mica dar mai mare ca 3
                                # regex10 = re.search('(.*)(\[)(.*)(\])',z)
                                # if regex10 != None:
                                #     regex11 = re.search('(.*)(.*?\s+(\S+))', regex3.group(1))
                                #     write_cell(1, 0, 0, regex11.group(1), table)
                                # else:
                                #     regex9 = re.search('(^[a-z]\w+)(.*)(.*?\s+(\S+))', z)
                                #     if regex9 != None:
                                #         write_cell(1, 0, 0, regex9.group(1), table)


                                regex9 = re.search('(^[a-z]\w+)(.*)(.*?\s+(\S+))', variable_definition)
                                if regex9 != None:
                                    regex10 = re.search('(.*)(\[)(.*)(\])', variable_definition)
                                    if regex10 != None:
                                        regex11 = re.search('(.*)(.*?\s+(\S+))', regex10.group(1))
                                        write_cell(1, 0, 0, regex11.group(1), table)
                                    else:
                                        write_cell(1, 0, 0, regex9.group(1), table)

                            elif (len(variable_definition.split())) == 2:
                                    write_cell(1, 0, 0, variable_definition.split()[0], table)

                             #lungime mai mica de 4
                            else:
                                regex1 = re.search('(^[A-Z])', variable_definition)
                                if regex1 != None:
                                    regex7 = re.search('(^[A-Z]\w+)(.*)(.*?\s+(\S+))', variable_definition)
                                    if regex7 != None:
                                        write_cell(1, 0, 0, regex7.group(2), table)
                                else:
                                    regex70 = re.search('(.*)(.*?\s+(\S+))', variable_definition)
                                    if regex70 != None:
                                        write_cell(1, 0, 0, regex70.group(1), table)

                            write_cell(0, 1, 1, "Value", table)
                            regex01 = re.search('(=)', variable_definition)
                            if regex01 != None:
                                regex02 = re.search('(=)(.*)', variable_definition)
                                if regex02 != None:
                                    write_cell(1, 1, 2, regex02.group(2), table)
                            else:
                                write_cell(1, 1, 2, "NA", table)
                            # ========================================================

                            new_entry(2, "Description", table)
                            write_cell(3, 0, 2, variable_description, table)

                            new_entry(4, "Definition", table)
                            write_cell(5, 0, 2, variable_definition, table)

                            curr_raw = 6

                            # ===========================================================
                            # WRITE DIFFERENT PROPERTIES FOUND PARSING MAN3 FILES
                            # ===========================================================
                            # If properties found
                            if len(lines_to_add_in_table) != 0:

                                # For every property in appended list
                                for element in lines_to_add_in_table:

                                    # Write property name
                                    new_entry(curr_raw, element[0], table)
                                    curr_raw += 1

                                    # Write property value
                                    write_cell(curr_raw, 0, 2, element[1], table)
                                    curr_raw += 1
                            # ===========================================================

                            if 'Group' in file:
                                group_variables_list_per_file.append(variable_name)

        if group_variables_list_per_file != []:
            if file in group_statistics:
                group_statistics[file]['Variables'] = group_variables_list_per_file
            else:
                group_statistics[file] = {}
                group_statistics[file]['Variables'] = group_variables_list_per_file
    doc.add_paragraph(" ")
    doc.save(new_doc_path)


def define_parser():

    print('PARSING DEFINES')

    sort_list = []
    doc.add_heading("Macros", level=2)

    macro_name = ''
    for file in module_statistics:

        group_macros_list_per_file = []

        for x in module_statistics[file]:
            if x == 'macros':
                # for element and it's index in list
                # for variable, list_increment in zip( module_statistics[file][x], tdmq(module_statistics[file][x]) ):
                print('\n', "Parsing File:", file )
                for define, list_increment in zip( module_statistics[file][x],tqdm(module_statistics[file][x]) ) :

                    macro_definition = ''
                    macro_description = "NA"
                    macro_definition_filter_list = ''

                    # If there is variables in parsed list
                    if define != '-':

                        # if variable found is in fact a list
                        if isinstance(define, list):

                            # If a string far matching a variable has been found
                            # add conditions to filter a better way all data
                            if isinstance(define[0], str) and len(define[0]) > 5 and define[0] not in macro_definition_filter_list and '#define' in define[0]:
                                macro_definition = define[0]

                            # If a string far matching a variable description has been found
                            if isinstance(define[1], str) and len(define[1]) > 5 and define[0] not in macro_definition_filter_list:
                                macro_description = define[1]

                        else:
                            if define not in macro_definition_filter_list and '#define' in define:
                                macro_definition = define
                            else:
                                macro_definition = ''

                    # WRITE MACRO DEFINITION IN DOCUMENT AS HEADER
                    # ====================================================================

                    # sort them into a list
                    if macro_definition not in sort_list and macro_definition != '':

                        sort_list.append(macro_definition)

                        # Add heading
                        regex4 = re.search('(.*?\s+(\S+))', macro_definition)
                        if regex4 != None:
                            doc.add_heading(regex4.group(2), level=3)
                            macro_name = regex4.group(2)

                            table = doc.add_table(rows=6, cols=3)
                            table.style = 'Table Grid'

                            curr_raw = 0
                            write_cell(0, 0, 1, "Name", table)
                            write_cell(0, 2, 0, "Value", table)
                            regex1 = re.search('(\()', macro_definition)
                            if regex1 != None:
                                regex0 = re.search('(\+)',macro_definition)
                                if regex0 != None:
                                    regex5 = re.search('((.*?)\()(.*?\))',macro_definition)
                                    if regex5 != None:
                                        #regex6 = re.search('(.*?\s+(\S+))', regex5.group(2))
                                        #if regex6 != None:
                                            #write_cell(1, 0, 1, regex6.group(2), table)
                                        write_cell(1, 0, 1, regex4.group(2), table)
                                        write_cell(1, 2, 0, regex5.group(3), table)
                                else:
                                    #regex2 = re.search('(.*)(.*?\s+(\S+))((\()(.*))(\))', z)
                                    #regex2 = re.search('(.*)((\()(.*)(\)))', z)
                                    regex2 = re.search('((\()(.*))', macro_definition)
                                    if regex2 != None:
                                        # regex6 = re.search('(.*?\s+(\S+))', regex2.group(1))
                                        # if regex6 != None:
                                        #     write_cell(1, 0, 1, regex6.group(2), table)
                                        #     print("intra2")
                                        write_cell(1, 0, 1, regex4.group(2), table)
                                        write_cell(1, 2, 0, regex2.group(1), table)
                            elif len(macro_definition.split()) >= 3:
                                regex3 = re.search('(.*)(.*?\s+(\S+))', macro_definition)
                                if regex3 != None:
                                    #write_cell(1, 0, 1, regex3.group(1), table)
                                    write_cell(1, 0, 1, regex4.group(2), table)
                                    write_cell(1, 2, 0, regex3.group(2), table)
                            else:
                                #write_cell(1, 0, 1, z, table)
                                write_cell(1, 0, 1, regex4.group(2), table)
                                write_cell(1, 2, 0, "NA", table)

                            # WRITE MACRO DEFINITION AND DESCRIPTION IN DOCUMENT IN TABLE
                            # ====================================================================

                            new_entry(2, "Definition", table)
                            write_cell(3, 0, 2, macro_definition, table)

                            new_entry(4, "Description", table)
                            write_cell(5, 0, 2, macro_description, table)

                            if macro_name != '' and 'Group' in file:
                                group_macros_list_per_file.append(macro_name)

        # WRITE MACRO IN DICT IF MACRO HAS FOUND IN GROUP
        # ====================================================================
        if group_macros_list_per_file != []:
            if file in group_statistics:
                group_statistics[file]['Macros'] = group_macros_list_per_file
            else:
                group_statistics[file] = {}
                group_statistics[file]['Macros'] = group_macros_list_per_file
    doc.add_paragraph(" ")
    doc.save(new_doc_path)

def structure_parser():
    # if there are elements in structure statistics
    if bool(structure_statistics):

        print('PARSING STRUCTURES')

        # add heading in document
        doc.add_heading("Structures", level=2)

        for file in structure_statistics:
            # add structure title in document
            doc.add_heading(file.replace('.3', ''), level=3)

            # add table in document
            table = doc.add_table(rows=0, cols=3)
            table.style = 'Table Grid'

            # define current row
            current_row = 0

            # write definition in table
            table.add_row()
            new_entry(current_row, "Definition", table)
            current_row += 1

            table.add_row()
            write_cell(current_row, 0, 2, file.replace('.3', ''), table)
            current_row += 1

            for feature in structure_statistics[file]:

                # write feature in table
                # HERE ARE WRITTEN IN TABLE HEADINGS, LIKE FIELDS, DESCRIPTION, ETC.
                table.add_row()
                new_entry(current_row, feature, table)
                current_row += 1

                if feature == 'Description':

                    # [ struct description, [ Remarks, a struct remark ] ]
                    if isinstance(structure_statistics[file][feature][0], list):

                        for type in structure_statistics[file][feature][0]:

                            # [ HERE WE ARE, [ .... ] ]
                            if isinstance(type, str):
                                table.add_row()
                                write_cell(current_row, 0, 2, type, table)
                                current_row += 1

                            # [ struct description [ HERE WE ARE ] ]
                            if isinstance(type, list):
                                table.add_row()
                                new_entry(current_row, type[0], table)
                                current_row += 1

                                table.add_row()
                                write_cell(current_row, 0, 2, type[1], table)
                                current_row += 1
                    else:
                        table.add_row()
                        write_cell(current_row, 0, 2, structure_statistics[file][feature][0], table)
                        current_row += 1

                if feature == 'Fields':

                    # Write subrows Name and Description
                    table.add_row()
                    sub_entry(current_row, 2, "Name Description", "#C6D9F1", table)
                    b = table.cell(current_row, 1).merge(table.cell(current_row, 2))
                    current_row += 1

                    # Write fields
                    for type in structure_statistics[file][feature]:
                        if isinstance(type, list):
                            table.add_row()
                            write_cell(current_row, 0, 0, type[0], table)
                            write_cell(current_row, 1, 0, type[1], table)
                            b = table.cell(current_row, 1).merge(table.cell(current_row, 2))
                            current_row += 1

                        if isinstance(type, str):
                            table.add_row()
                            write_cell(current_row, 0, 0, type, table)
                            write_cell(current_row, 1, 0, '*', table)
                            b = table.cell(current_row, 1).merge(table.cell(current_row, 2))
                            current_row += 1


# Functions to be executed at the final of entire script
def post_main_functions():

    # ===================================================
    # WRITE IN DOCUMENT TABLE FROM GROUPS SECTION
    # ===================================================
    # if dictionary group statistics is not empty
    if bool(group_statistics):
        print('ADDING GROUPS')

        # Define how paragraph with groups table shall name
        group_paragraph_name = 'Groups'

        # Write heading in document
        # List all paragraphs in document
        paragraphs = list(doc.paragraphs)

        group_paragraph = paragraphs[0]

        # add paragraph in document
        for idx, paragraph in enumerate(paragraphs):
            if "Features" == paragraph.text:

                # add Title
                prior_paragraph = paragraphs[idx - 1].insert_paragraph_before(group_paragraph_name)
                prior_paragraph.style = 'Heading 2'

                # Retain group paragraph
                group_paragraph = prior_paragraph

        for file in group_statistics:

            # add table at final of document
            table_groups = doc.add_table(rows=0, cols=1)
            table_groups.style = 'Table Grid'

            # Do not allow table to autofit itself
            table_groups.autofit = False
            table_groups.allow_autofit = False

            current_row = 0

            table_groups.add_row()
            table_groups.cell(current_row, 0).text = 'Group Name'
            current_row += 1

            table_groups.add_row()
            table_groups.cell(current_row, 0).text = file.replace('.3', '')
            current_row += 1

            for feature in group_statistics[file]:
                table_groups.add_row()
                new_entry_simple_table(row=current_row,text= feature, table_obj=table_groups)
                current_row += 1

                for entity in group_statistics[file][feature]:
                    table_groups.add_row()
                    table_groups.cell(current_row, 0).text = entity
                    current_row += 1

            # move table after founded paragraph
            move_table_after(table_groups, group_paragraph)

    # ===================================================
    # ADD LAST HEADINGS IN DOCUMENT
    # ===================================================
    # add headings to final of document
    doc.add_heading("EEPROM", level=1)
    doc.add_paragraph(" ")
    doc.add_heading("Configuration", level=1)
    doc.add_paragraph(" ")
    doc.add_heading("Compilation Options", level=1)
    doc.add_paragraph(" ")
    doc.save(new_doc_path)

def add_graphics_in_word_document():

    print("ADDING MORITZ GRAPHICS")

    # Call dispatch application for Word
    word = win32.gencache.EnsureDispatch('Word.Application')

    # Assign application path and set it active
    docx = word.Documents.Open(new_doc_path)
    word.Visible = False
    docx.Activate()

    list_of_functions_with_graphics = []
    # for variable, list_increment in zip( module_statistics[file][x], tdmq(module_statistics[file][x]) ):
    for function_name, list_increment in zip( list_of_functions, tqdm(list_of_functions) ):

        # get a list of all files from directory
        list_with_all_files_in_directory = os.listdir(path_to_moritz_folder_html)

        for file in list_with_all_files_in_directory:
            if 'dot_' in file and function_name in file and file.endswith('.png'):

                path_to_graphic_function = os.path.join(path_to_moritz_folder_html, file)

                # convert found png to bmp
                path_to_graphic_function_bmp = convert_png_files_to_bmp_files(file_to_convert=path_to_graphic_function)

                # if exists, for every paragraph in document
                for paragraph in docx.Paragraphs:

                    # If found file name is in word
                    if 'dot_' + function_name + '.png'  in paragraph.Range.Text:

                        # delete written name
                        paragraph.Range.Text = ''

                        # add object in word
                        docx.InlineShapes.AddOLEObject(FileName=path_to_graphic_function_bmp, Range=paragraph.Range )

                        # if needed, here are some configurations you can add
                        # paragraph.Range.InsertAfter("\r\n")
                        # docx.OLEObjects.Add(ClassType=None, Filename=path_to_graphic_function, Link=False, DisplayAsIcon=True,IconFileName=wordpath, IconIndex=0, IconLabel=Final.docx Left = dest_cell.Left, Top = dest_cell.Top, Width = 50, Height = 50)

                        list_of_functions_with_graphics.append(function_name)

    # Print a message on screen if function has no graphic generated
    for function in list_of_functions:
        if function not in list_of_functions_with_graphics:
            print(function + ' function has no graphic generated from Moritz tool! ')

    # because this word document is protected, it is needed to accept all changes to it
    word.ActiveDocument.Revisions.AcceptAll()

    # Delete all comments
    if word.ActiveDocument.Comments.Count >= 1:
        word.ActiveDocument.DeleteAllComments()

    # Save document and quit
    word.ActiveDocument.Save()
    docx.Application.Quit(-1)


if __name__ == "__main__":
    main()
    type_of_data = yaml_config_file['Data entity to parse']
    print('Parsing data:', type_of_data)
    if type_of_data == "functions":
        function_parser()
    if type_of_data == "variables":
        variable_parser()
    if type_of_data == "defines":
        define_parser()
    if type_of_data == "structures":
        structure_parser()
    if type_of_data == "all":
        function_parser()
        variable_parser()
        define_parser()
        structure_parser()
    post_main_functions()
    add_graphics_in_word_document()

