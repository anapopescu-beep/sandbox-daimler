"""Delivery_test_summary_report.py: Used to generate summary report based on excel ( RA, etc. ) as input
Parameters has been given from argparse command line interface.
As input, it is needed an path where tests can be found and to template
As output, it is needed a path to generate an docx file based upon given template
"""

__author__ = 'Ardeleanu Lucian'
__copyright__ = "Copyright 2021, Autoliv Electronic"
__version__ = "1.0"
__email__ = 'lucian.ardeleanu@autoliv.com'
__status__ = 'Released'

# ============ IMPORTS ====================
import xlrd
from docx import Document
from copy import deepcopy
import re
import argparse
try:
    import win32com.client
except:
    pass


# ===============================================================
# FUNCTION USED TO UPDATE THE TABLE OF CONTENT FOR DOCX FILES
# INPUT ARGUMENT: PATH TO DOCUMENT
# OUTPUT ARGUMENT: NONE
# ===============================================================
def update_toc(docx_file):
    word = win32com.client.DispatchEx("Word.Application")
    doc = word.Documents.Open(docx_file)
    toc_count = doc.TablesOfContents.Count
    if toc_count != 0:
        doc.TablesOfContents(1).Update()
        print('TOC should have been updated.')
    else:
        print('TOC has not been updated for sure...')
    doc.Close(SaveChanges=True)
    word.Quit()


# ===============================================================
# FUNCTION USED TO DELETE A PARAGRAPH
# INPUT ARGUMENT: OBJECT-BASED PARAGRAPH
# OUTPUT ARGUMENT: NONE
# ===============================================================
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


# ===============================================================
# FUNCTION USED TO EXTRACT VERSION NUMBER FROM AN STRING
# INPUT ARGUMENT: STRING
# OUTPUT ARGUMENT: STRING
# ===============================================================
def version_parser(v):
    versionPattern = r'\d+(=?\.(\d+(=?\.(\d+)*)*)*)*'
    regexMatcher = re.compile(versionPattern)
    return regexMatcher.search(v).group(0)


# ===============================================================
# FUNCTION USED TO REMOVE A ROW FROM A GIVEN TABLE
# INPUT ARGUMENT: OBJECT-BASED TABLE, OBJECT-BASED ROW
# OUTPUT ARGUMENT: NONE
# ===============================================================
def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


# ===============================================================
# FUNCTION USED TO INSERT PARAGRAPH AFTER SPECIFIC PARAGRAPHS
# INPUT ARGUMENT: LIST OF OBJECT-BASED PARAGRAPHS, INTEGER: INDEX,
#                 STRING: TEXT TO INSERT
# OUTPUT ARGUMENT: IT WILL INSERT A PARAGRAPH INSIDE DOCUMENT
# ===============================================================
def insert_paragraph_after(paragraphs, idx, text=None):
    next_paragraph_idx = idx + 1
    if idx == len(paragraphs):
        return document.add_paragraph(text)
    next_paragraph = paragraphs[next_paragraph_idx]
    return next_paragraph.insert_paragraph_before(text)


# ===============================================================
# FUNCTION USED TO OPEN EXCEL FILE AND EXTRACT DATA FROM IT
# INPUT ARGUMENT: STRING: PATH TO EXCEL FILE TO PROCESS
# OUTPUT ARGUMENT: LIST OF STRINGS WITH DATA FROM EXCEL
# ===============================================================
def extract_data_from_excel(file_to_process):

    # open workbook
    wb = xlrd.open_workbook(file_to_process)
    sheets = wb.sheet_names()

    # open sheet
    # check which file is needed for opening
    for sheet_name in sheets:
        if ( 'Delivery Test list' in sheet_name.lower() or 'delivery test list' in sheet_name.lower() or 'Delivery Test List' in sheet_name.lower() ):
            ws = wb[sheet_name]
            break

    # Get total number of rows and columns
    rows = ws.nrows

    # We don't actually need all columns for document
    #columns = ws.ncols

    # append document data in list
    list_of_global_data_from_excel =[]

    # Initialise first list that will contain only tests from xls
    list_with_test_cases = []
    for rows_counter in range(2, rows):

        # extract cell value
        cell_value = ws.cell_value(rows_counter, 1)

        # if cell is a kind of title
        if ( cell_value != '')  and ( cell_value != None ) and ('DELIV' not in cell_value) and ( len(version_parser(cell_value)) > 0 ) :
            # append only title
            list_of_global_data_from_excel.append(cell_value)

            # Initialise list with test cases again
            list_with_test_cases = []

        else:

            # if cell is a test case
            list_with_data_per_row = ['', '', '']
            if 'DELIV' in cell_value:
                # Extract test case data ( TEST ID, TESTER, STATUS )
                list_with_data_per_row[0] =  ws.cell_value(rows_counter, 1)
                list_with_data_per_row[1] =  ws.cell_value(rows_counter, 3)
                list_with_data_per_row[2] =  ws.cell_value(rows_counter, 7)

            # append test case in global list
            if ( list_with_data_per_row != ['', '', ''] ):
                list_with_test_cases.append(list_with_data_per_row)

        if list_with_test_cases not in list_of_global_data_from_excel and list_with_test_cases != []:
            # append in global list test cases
            list_of_global_data_from_excel.append(list_with_test_cases)


    return list_of_global_data_from_excel


# ===============================================================
# MAIN FUNCTION OF SCRIPT
# INPUT ARGUMENT: NONE
# OUTPUT ARGUMENT: NONE
# ===============================================================
def main():

    parser = argparse.ArgumentParser()
    parser.add_argument('-i', '--input-path-ra-reports', help="Path to Delivery Tests Follow-up xls", required=True)
    parser.add_argument('-l', '--input-path-template-docx', help="Path to Template DOCX ( Summary Report )  with filename and extension", required=True)
    parser.add_argument('-o', '--output-path-generated-docx', help="Path to generate DOCX Summary Report with filename and extension",required=True)
    args = parser.parse_args()

    path_to_test_reports = args.input_path_ra_reports
    path_to_template_docx_file = args.input_path_template_docx
    path_to_generated_docx_file = args.output_path_generated_docx

    # -i="s:\Tests\Tests_Deliveries\Automatic_Tests\SW Delivery Tests follow-up.xls" -l="Audi_Tranche_6_Integration_Tests_Summary_Report_template.docx" -o="Test.docx"

    # path_to_test_reports = "S:\Tests\Tests_Integration\Automatic_Tests\Test_Reports"
    # path_to_template_docx_file = "Audi_Tranche_6_Integration_Tests_Summary_Report_template.docx"
    # path_to_generated_docx_file = "Test.docx"

    # extract in list content of ra file
    list_of_xls_content = extract_data_from_excel( file_to_process= path_to_test_reports)

    if len(list_of_xls_content) == 0:
        print('DATA CANNOT BE EXTRACTED FROM GIVEN INPUT, NO DOCUMENT WILL BE GENERATED')
        return 0
    else:
        print('DATA HAS BEEN EXTRACTED SUCCESFULLY!')

    # ====================================================================
    # WRITE EXTRACTED DATA IN DOCX DOCUMENT
    # ====================================================================

    # Get an instance of document
    document = Document(path_to_template_docx_file)

    document.save(path_to_generated_docx_file)

    # ===================================
    # ADD HEADINGS TO DOCUMENT
    # ===================================

    counter_for_heading_styles = 3

    for element in list_of_xls_content:
        # Open document
        document = Document(path_to_generated_docx_file)

        # List all paragraphs in document
        paragraphs = list(document.paragraphs)

        # Iterare and index all paragraphs
        for idx, paragraph in enumerate(paragraphs):

            # if element is string ( if element is a title )
            if isinstance(element, str):

                # if found a specific paragraph
                if 'Global Summary' in paragraph.text:

                    # add Title
                    prior_paragraph = paragraphs[idx - 1].insert_paragraph_before(element[ len(version_parser(element)) : ])
                    prior_paragraph.style = 'Heading ' + str(counter_for_heading_styles)

                    counter_for_heading_styles += 1

            else:
                counter_for_heading_styles = 3

        document.save(path_to_generated_docx_file)

    # ============================================================
    # GET LAST HEADINGS FROM XLS LIST (WHERE DATA START TO BEGIN)
    # ============================================================

    # init a new list with headings
    list_with_last_headings = []

    # init a list only with test cases
    list_with_tests_only = []

    # init the lists of tests counters
    list_of_OK_tests_counter = []
    list_of_NOK_tests_counter = []
    total_number_of_OK_tests = 0
    total_number_of_NOK_tests = 0

    # for every element in global list from xls
    for index_element in range(0,len(list_of_xls_content)):

        # if found element is a list
        if isinstance(list_of_xls_content[index_element], list):
            # init counters for OK and NOK tests
            OK_tests_counter = 0
            NOK_tests_counter = 0

            # get the number of OK and NOK tests
            for index_line in range(0,len(list_of_xls_content[index_element])):

                if list_of_xls_content[index_element][index_line][2] == 'OK':
                    OK_tests_counter += 1

                if list_of_xls_content[index_element][index_line][2] == 'NOK':
                    NOK_tests_counter += 1

                # fill empty tests status with NOTRUN
                if list_of_xls_content[index_element][index_line][2] == '':
                    list_of_xls_content[index_element][index_line][2] = 'NOTRUN'

            # add current counters to the total number
            total_number_of_OK_tests += OK_tests_counter
            total_number_of_NOK_tests += NOK_tests_counter

            # append the counter to the list of counters
            list_of_OK_tests_counter.append(OK_tests_counter)
            list_of_NOK_tests_counter.append(NOK_tests_counter)

            # append element in list with test cases only
            list_with_tests_only.append(list_of_xls_content[index_element])

            # if previous element is a string
            if isinstance(list_of_xls_content[index_element - 1 ], str):

                # Retain string in list WITHOUT ANY VERSION IN IT
                list_with_last_headings.append(list_of_xls_content[index_element - 1][ len(version_parser(list_of_xls_content[index_element - 1])) : ] )

    # ===================================
    # ADD TABLES TO DOCUMENT
    # ===================================

    # Open document
    document = Document(path_to_generated_docx_file)

    # List all paragraphs in document
    paragraphs = list(document.paragraphs)

    # count where are last headings
    counter_for_last_headings = 0

    # Variable used to count how many tables has been added
    counter_for_added_tables = 0

    # Iterare and index all paragraphs
    for idx, paragraph in enumerate(paragraphs):

        # Select table needed to be copied
        template = document.tables[4]

        # Here we do the copy of the table
        tbl = template._tbl
        new_tbl = deepcopy(tbl)

        # if found a specific paragraph
        if list_with_last_headings[counter_for_last_headings] in paragraph.text:

            # ========== ADD TABLE ======================

            # Insert table
            insert_paragraph_after(paragraphs, idx, paragraph._p.addnext(new_tbl))
            counter_for_added_tables += 1

            # ========== ADD STATISTICS ======================

            # Add an empty paragraph
            prior_paragraph = insert_paragraph_after(paragraphs, idx, "\r\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run('Number of tests: '
                                          + str(len(list_with_tests_only[counter_for_last_headings])))
            run.bold = True
            run.italic = True

            # Add a new line to document
            prior_paragraph.add_run("\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run('Number of positives tests: '
                                          + str(list_of_OK_tests_counter[counter_for_last_headings]))
            run.bold = True
            run.italic = True

            # Add a new line to document
            prior_paragraph.add_run("\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run('Number of negatives tests: '
                                          + str(list_of_NOK_tests_counter[counter_for_last_headings]))
            run.bold = True
            run.italic = True

            # Add a new line to document
            prior_paragraph.add_run("\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run('Number of not run tests: '
                                          + str(len(list_with_tests_only[counter_for_last_headings])
                                                - list_of_OK_tests_counter[counter_for_last_headings]
                                                - list_of_NOK_tests_counter[counter_for_last_headings]))
            run.bold = True
            run.italic = True

            # count in list with headings
            if counter_for_last_headings <= len(list_with_last_headings):
                counter_for_last_headings += 1

        # if list of heading has arrived to final
        if counter_for_last_headings >= len(list_with_last_headings):
            break


    # DELETE SECTION WITH "Tests Performed ( Including table, we copied it in upper code )"
    delete_paragraph(document.tables[4])
    for idx, paragraph in enumerate(paragraphs):
        if "Tests Performed" in paragraph.text:
            delete_paragraph(paragraph)


    # ADD ADD TABLES ROWS AND FILL THEM WITH DATA

    for ( table,set_of_tests) in zip(document.tables[4: 4 + counter_for_added_tables], list_with_tests_only):

        # ==================================
        # ADD/DELETE TABLE ROWS
        # ==================================

        # COUNT HOW MANY ROWS ARE IN TABLE
        number_of_table_rows = 0
        for row in table.rows:
            # Get number of maximum rows
            number_of_table_rows += 1

        # GET DIFFERENCE OF ROWS AND ADD IF NEEDED MORE
        difference_of_rows = (number_of_table_rows - 1) - len(set_of_tests)

        # If there are more rows than needed
        if (difference_of_rows > 0):
            # Delete exceeded rows
            for counter_for_difference_rows in range(0, difference_of_rows):
                row = table.rows[difference_of_rows - counter_for_difference_rows]
                remove_row(table, row)
        else:
            # Else, add rows to table
            for counter_for_difference_rows in range(0, -difference_of_rows):
                table.add_row()

        # ============================
        # ADD DATA IN TABLE
        # ============================

        # Do not allow table to autofit itself
        table.autofit = False
        table.allow_autofit = False

        for index_test in range(0, len(set_of_tests) ):

            # ======== WRITE IN FIRST CELL ==========
            # ACCESS CELL FROM TABLE
            cell = table.cell(index_test + 1, 0) # first parameter is row, second is column

            # access paragraph from cell
            paragraph_from_cell = cell.paragraphs[0]

            # add text to it and bold it
            run = paragraph_from_cell.add_run( set_of_tests[index_test][0].upper() )
            run.bold = True

            # ======== WRITE IN SECOND CELL ==========
            # ACCESS CELL FROM TABLE
            cell = table.cell(index_test + 1, 1) # first parameter is row, second is column

            # access paragraph from cell
            paragraph_from_cell = cell.paragraphs[0]

            # add text to it and bold it
            run = paragraph_from_cell.add_run( set_of_tests[index_test][1].upper() )
            run.bold = True

            # ======== WRITE IN THIRD CELL ==========
            # ACCESS CELL FROM TABLE
            cell = table.cell(index_test + 1, 2) # first parameter is row, second is column

            # access paragraph from cell
            paragraph_from_cell = cell.paragraphs[0]

            # add text to it and bold it
            run = paragraph_from_cell.add_run( set_of_tests[index_test][2].upper() )
            run.bold = True

    counter_needed_to_sum_all_tests = 0
    for set_of_tests in list_with_tests_only:
        counter_needed_to_sum_all_tests += len(set_of_tests)

    # ADD TOTAL STATISTICS TO DOCUMENT
    # Iterate and index all paragraphs
    for idx, paragraph in enumerate(paragraphs):
        # if found a specific paragraph
        if 'Total number of tests:' in paragraph.text:
            # Add new text to it
            run = paragraph.add_run(str(counter_needed_to_sum_all_tests))
            run.bold = True
            run.italic = True

        if 'Total number of positives tests:' in paragraph.text:
            # Add new text to it
            run = paragraph.add_run(str(total_number_of_OK_tests))
            run.bold = True
            run.italic = True

        if 'Total number of negatives tests:' in paragraph.text:
            # Add new text to it
            run = paragraph.add_run(str(total_number_of_NOK_tests))
            run.bold = True
            run.italic = True

        if 'Total number of not run tests:' in paragraph.text:
            # Add new text to it
            run = paragraph.add_run(str(counter_needed_to_sum_all_tests - total_number_of_OK_tests - total_number_of_NOK_tests))
            run.bold = True
            run.italic = True


    document.save(path_to_generated_docx_file)
    try:
        update_toc(path_to_generated_docx_file)
    except:
        print("Table of content not updated!!!")


if __name__ == "__main__":
    main()