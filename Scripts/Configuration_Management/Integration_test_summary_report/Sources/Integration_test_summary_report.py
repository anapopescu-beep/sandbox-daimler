"""Integration_test_summary_report.py: Used to generate summary report based on excel ( RA, etc. ) as input
Parameters has been given from argparse command line interface.
As input, it is needed an path where tests can be found and to template
As output, it is needed a path to generate an docx file based upon given template
"""

__author__ = 'Ardeleanu Lucian'
__copyright__ = "Copyright 2021, Autoliv Electronic"
__version__ = "1.0"
__email__ = 'lucian.ardeleanu@autoliv.com'
__status__ = 'Released'

# ============ IMPORTS ====================
import xlrd
import os
from docx import Document
from copy import deepcopy
import argparse
try:
    import win32com.client
except:
    pass

# ===============================================================
# FUNCTION USED TO UPDATE THE TABLE OF CONTENT FOR DOCX FILES
# INPUT ARGUMENT: PATH TO DOCUMENT
# OUTPUT ARGUMENT: NONE
# ===============================================================
def update_toc(docx_file):
    word = win32com.client.DispatchEx("Word.Application")
    doc = word.Documents.Open(docx_file)
    toc_count = doc.TablesOfContents.Count
    if toc_count != 0:
        doc.TablesOfContents(1).Update()
        print('TOC should have been updated.')
    else:
        print('TOC has not been updated for sure...')
    doc.Close(SaveChanges=True)
    word.Quit()


# ===============================================================
# FUNCTION USED TO REMOVE A ROW FROM A GIVEN TABLE
# INPUT ARGUMENT: OBJECT-BASED TABLE, OBJECT-BASED ROW
# OUTPUT ARGUMENT: NONE
# ===============================================================
def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


# ===============================================================
# Move the value <value> from the key <src_key> to the key <dest_key>
# Raises ValueError if the value was not found in the list of the keywork
# ===============================================================
def moveBetweenKeys(dictionary, src_key, dest_key, value):
    src_list = dictionary[src_key]
    dest_list = dictionary[dest_key]

    if value not in src_list:
        raise ValueError("'%s' was not found in the list '%s'" % (value, src_key))

    src_list.remove(value)
    dest_list.append(value)


# ===============================================================
# FUNCTION USED TO INSERT PARAGRAPH AFTER SPECIFIC PARAGRAPHS
# INPUT ARGUMENT: LIST OF OBJECT-BASED PARAGRAPHS, INTEGER: INDEX,
#                 STRING: TEXT TO INSERT
# OUTPUT ARGUMENT: IT WILL INSERT A PARAGRAPH INSIDE DOCUMENT
# ===============================================================
def insert_paragraph_after(paragraphs, idx, text=None):
    next_paragraph_idx = idx + 1
    if idx == len(paragraphs):
        return document.add_paragraph(text)
    next_paragraph = paragraphs[next_paragraph_idx]
    return next_paragraph.insert_paragraph_before(text)


# ===============================================================
# FUNCTION USED TO OPEN EXCEL FILE AND EXTRACT DATA FROM IT
# INPUT ARGUMENT: STRING: PATH TO EXCEL FILE TO PROCESS
# OUTPUT ARGUMENT: LIST OF STRINGS WITH DATA FROM EXCEL
# ===============================================================
def extract_data_from_excel(file_to_process):
    # Define List in witch data will be stored
    raw_list = []

    # open workbook
    wb = xlrd.open_workbook(file_to_process)
    sheets = wb.sheet_names()

    # open sheet
    # check which file is needed for opening

    for sheet_name in sheets:
        if ( 'SRM' in sheet_name.lower() or 'srm' in sheet_name.lower() ):
            ws = wb[sheet_name]
            break

    # Get total number of rows and columns
    rows = ws.nrows

    # We don't actually need all columns for document
    #columns = ws.ncols
    columns = 7

    # append document data in list
    for i in range(3, rows):
        for j in range(6, columns):

            # Extract cell value in variable
            cell_value = str(ws.cell_value(i, j)).lower()

            # If cell value is not empty
            if cell_value != '':

                # Remove all spaces
                cell_value = cell_value.replace(" ", "")

                # Replace , with ;
                if ',' in cell_value:
                    cell_value = cell_value.replace("," , ";")

                # Split after every ; and append in list
                if ";" in cell_value:
                    for element in cell_value.split(';'):
                        if element not in raw_list:
                            raw_list.append(element)
                else:
                    if cell_value not in raw_list:
                        raw_list.append(cell_value)
    return raw_list


# ===============================================================
# FUNCTION USED TO OPEN XML FILES AND EXTRACT TESTS STATUS
# INPUT ARGUMENT:STRING: PATH TO XML FILE TO PROCESS
# OUTPUT ARGUMENT:STRING: STATUS OF THE TEST
# ===============================================================
def extract_data_from_xml(file_to_process):
    # status is read flag
    status_is_read = 0
    # initialize status as OK
    status_return = "NOK"
    # Open file for reading text.
    with open(file_to_process, 'r') as myfile:
        # For each line of text
        for line in myfile:
            # verify if the line contain "Result Test:" and if the line not contain "Passed".
            if "Result Test:" in line:
                status_is_read = 1
                if 'PASSED' in line:
                    # actualize the status as "NOK"
                    status_return = "OK"
            if status_is_read:
                break
    return status_return


# ===============================================================
# MAIN FUNCTION OF SCRIPT
# INPUT ARGUMENT: NONE
# OUTPUT ARGUMENT: NONE
# ===============================================================
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('-i', '--input-path-ra-reports', help="Path to Test Reports where RA files are", required=True)
    parser.add_argument('-l', '--input-path-template-docx', help="Path to Template DOCX ( Summary Report )  with filename and extension", required=True)
    parser.add_argument('-o', '--output-path-generated-docx', help="Path to generate DOCX Summary Report with filename and extension",required=True)
    args = parser.parse_args()

    path_to_test_reports = args.input_path_ra_reports
    path_to_template_docx_file = args.input_path_template_docx
    path_to_generated_docx_file = args.output_path_generated_docx

    # -i="S:\Tests\Tests_Integration\Automatic_Tests\Test_Reports" -l="Audi_Tranche_6_Integration_Tests_Summary_Report_template.docx" -o="Test.docx"

    # path_to_test_reports = "S:\Tests\Tests_Integration\Automatic_Tests\Test_Reports"
    # path_to_template_docx_file = "Audi_Tranche_6_Integration_Tests_Summary_Report_template.docx"
    # path_to_generated_docx_file = "Test.docx"

    dict_with_all_data_from_ra = {}
    for root, dirs, files in os.walk(path_to_test_reports):
        for file in files:
            if (file.endswith(".xls") and 'RA' in  file):

                # extract module name from file name
                module_name =  file[file.find("RA_") + len("RA_") : file.find(".xls") ]

                # Get full path to ra file
                path_to_ra_file = os.path.join(root, file)

                # extract in list content of ra file
                list_of_ra_content = extract_data_from_excel( file_to_process= path_to_ra_file)

                # append data in dictionary
                dict_with_all_data_from_ra[module_name] = list_of_ra_content

    if len(dict_with_all_data_from_ra) == 0:
        print('DATA CANNOT BE EXTRACTED FROM GIVEN INPUT, NO DOCUMENT WILL BE GENERATED')
        return 0
    else:
        print('DATA HAS BEEN EXTRACTED SUCCESFULLY!')

    # # check if tests are in corect modules
    # list_of_data_to_be_modified_global = []
    #
    # for key, value in dict_with_all_data_from_ra.items():
    #     for test_name in value:
    #         # Initialise a list with data to be modified in dictionary
    #         # Elements in list are: source_key in dictionary, destination key in dictionary and test name to be moved from one key to another
    #         list_of_data_to_be_modified = ['','','']
    #
    #         # if module name not in test name
    #         if key.lower() not in test_name:
    #
    #             # if ecu is not in test name ( because this are tests of SWarchitecture module )
    #             if "_ecu_" not in test_name:
    #
    #                 # Extract module name from file name
    #                 module_name_of_founded_file = test_name[ test_name.find('arch_test_') + len( 'arch_test_')  :   len('arch_test_') + test_name.find( '_') ]
    #                 module_name_of_founded_file = module_name_of_founded_file[ 0 : module_name_of_founded_file.find("_")]
    #                 module_name_of_founded_file = module_name_of_founded_file.upper()
    #
    #                 # found something to be modified, so let's append it into list
    #                 list_of_data_to_be_modified = [key, module_name_of_founded_file, test_name]
    #
    #         # if data to be modified has not been already in big list, let's append it
    #         if list_of_data_to_be_modified not in list_of_data_to_be_modified_global and list_of_data_to_be_modified != ['','','']:
    #             list_of_data_to_be_modified_global.append(list_of_data_to_be_modified)
    #
    # # Iterare throu list of data to be modified
    # for element in list_of_data_to_be_modified_global:
    #     if element[0] == "MMG":
    #         list_of_data_to_be_modified_global.remove(element)
    #     else:
    #         # print(element)
    #
    #         # if module name has not been in dictionary, added it
    #         if element[1] not in dict_with_all_data_from_ra:
    #             dict_with_all_data_from_ra[element[1]] = []
    #
    #         # move elements from module to another
    #         moveBetweenKeys(dictionary= dict_with_all_data_from_ra, src_key=element[0], dest_key=element[1], value=element[2])

    # For debug purpose
    # for key, value in dict_with_all_data_from_ra.items():
    #     print(len(value), key, ' : ', value )

    # ====================================================================
    # WRITE EXTRACTED DATA IN DOCX DOCUMENT
    # ====================================================================

    # Get an instance of document
    document = Document(path_to_template_docx_file)
    document.save(path_to_generated_docx_file)

    # Get all values from dictionary in a list
    list_of_values_from_dictionary = list(dict_with_all_data_from_ra.values())
    list_of_keys_from_dictionary = list(dict_with_all_data_from_ra.keys())

    # Initialize test counters
    number_of_OK_tests = []
    number_of_NOK_tests = []
    number_of_total_OK_tests = 0
    number_of_total_NOK_tests = 0
    number_of_OK_tests.append(0)
    number_of_NOK_tests.append(0)

    # print(list_of_values_from_dictionary)

    # ===================================
    # ADD HEADINGS AND TABLES TO DOCUMENT
    # ===================================

    for counter in range(0,len(dict_with_all_data_from_ra) - 1):
        # Open document
        document = Document(path_to_generated_docx_file)

        # Select table needed to be copied
        template = document.tables[4]

        # Here we do the copy of the table
        tbl = template._tbl
        new_tbl = deepcopy(tbl)

        # List all paragraphs in document
        paragraphs = list(document.paragraphs)

        # Iterare and index all paragraphs
        for idx, paragraph in enumerate(paragraphs):
            # if found a specific paragraph
            if 'Global Summary' in  paragraph.text :
                # add Title
                prior_paragraph = paragraphs[idx - 1].insert_paragraph_before("Tests Performed – "   )
                prior_paragraph.style = 'Heading 3'

                # add Table
                prior_paragraph = paragraph.insert_paragraph_before(paragraphs[idx - 1]._p.addnext(new_tbl))

        document.save(path_to_generated_docx_file)

    # =================================
    # ADD DATA IN INSERTED PARAGRAPHS
    # =================================

    # Open document
    document = Document(path_to_generated_docx_file)

    # List all paragraphs in document
    paragraphs = list(document.paragraphs)

    counter_for_list_of_values_from_dict = 0

    for table in document.tables[4: 4 + len(list_of_keys_from_dictionary) ]:

        # COUNT HOW MANY ROWS ARE IN TABLE
        number_of_table_rows = 0
        for row in table.rows:
            # Get number of maximum rows
            number_of_table_rows += 1

        # GET DIFFERENCE OF ROWS AND ADD IF NEEDED MORE
        difference_of_rows = ( number_of_table_rows - 1 ) - len(list_of_values_from_dictionary[counter_for_list_of_values_from_dict])


        # If there are more rows than needed
        if ( difference_of_rows > 0):
            # Delete exceeded rows
            for counter_for_difference_rows in range(0, difference_of_rows ):
                row = table.rows[difference_of_rows - counter_for_difference_rows ]
                remove_row(table, row)
        else:
            # Else, add rows to table
            for counter_for_difference_rows in range(0, -difference_of_rows):
                table.add_row()

        # Do not allow table to autofit itself
        table.autofit = False
        table.allow_autofit = False

        for index_test in range(0, len(list_of_values_from_dictionary[counter_for_list_of_values_from_dict]) ):

            # ======== WRITE IN FIRST CELL ==========
            # ACCESS CELL FROM TABLE
            cell = table.cell(index_test + 1, 0) # first parameter is row, second is column

            # access paragraph from cell
            paragraph_from_cell = cell.paragraphs[0]

            # add text to it and bold it
            run = paragraph_from_cell.add_run( list_of_values_from_dictionary[counter_for_list_of_values_from_dict][index_test] )
            run.bold = True

            # ======== WRITE IN SECOND CELL ==========
            # ACCESS CELL FROM TABLE
            cell = table.cell(index_test + 1, 1) # first parameter is row, second is column

            # access paragraph from cell
            paragraph_from_cell = cell.paragraphs[0]

            # add text to it and bold it
            run = paragraph_from_cell.add_run("SBE")
            run.bold = True

            # ======== WRITE IN THIRD CELL ==========
            # ACCESS CELL FROM TABLE
            cell = table.cell(index_test + 1, 2) # first parameter is row, second is column

            # access paragraph from cell
            paragraph_from_cell = cell.paragraphs[0]

            # add text to it and bold it
            test_status = "NOTRUN"  # initialize test_status as "NOTRUN"
            test_reports_dir_list = os.listdir(path_to_test_reports)    # list of directors in test reports folder
            for root, dirs, files in os.walk(path_to_test_reports):
                for file in files:
                    # verify if the current xml file is important
                    if file.lower().endswith(
                            list_of_values_from_dictionary[counter_for_list_of_values_from_dict][index_test] + ".xml"):
                        # the function extract_data_from_xml return the status of the current xml file, "OK" or "NOK"
                        test_status = extract_data_from_xml(root + '\\' + file)


            run = paragraph_from_cell.add_run(test_status)
            run.bold = True

            if test_status == "OK":
                number_of_OK_tests[counter_for_list_of_values_from_dict] += 1
            if test_status == "NOK":
                number_of_NOK_tests[counter_for_list_of_values_from_dict] += 1

        number_of_total_OK_tests = number_of_total_OK_tests + number_of_OK_tests[counter_for_list_of_values_from_dict]
        number_of_total_NOK_tests = number_of_total_NOK_tests + number_of_NOK_tests[counter_for_list_of_values_from_dict]

        number_of_OK_tests.append(0)
        number_of_NOK_tests.append(0)

        counter_for_list_of_values_from_dict += 1


    # Add modules names to modules tested list
    counter_for_keys_from_dict = 0
    for idx, paragraph in enumerate(paragraphs):
        # if found a specific paragraph
        if 'Tests Performed –' in paragraph.text:
            # Add dictionary key ( module names ) to previously added headings
            run = paragraph.add_run(list_of_keys_from_dictionary[counter_for_keys_from_dict])
            counter_for_keys_from_dict += 1

    # LIST ALL STYLES OF DOCUMENT BEFORE USING IT
    # for style in document.styles:
    #      print("style.name == %s" % style.name)

    # Add modules names to modules tested list
    for idx, paragraph in enumerate(paragraphs):
        # if found a specific paragraph
        if 'Priority #1 modules tested:' in paragraph.text:
            # add an empty paragraph after upper paragraph
            prior_paragraph = insert_paragraph_after(paragraphs, idx, "\n")
            prior_paragraph.style = 'Normal'
            for key, value in dict_with_all_data_from_ra.items():
                run = prior_paragraph.add_run(" - " + key + "\n")


    # Count in list of values from dictionary
    counter_for_list_of_values_from_dictionary = 0
    counter_of_all_tests = 0

    # Iterare and index all paragraphs
    for idx, paragraph in enumerate(paragraphs):
        # if found a specific paragraph
        if 'Tests Performed' in paragraph.text:
            # Add an empty paragraph
            prior_paragraph = insert_paragraph_after(paragraphs, idx + 1, "\r\n")

            # Add paragraphs with statistics

            # Add paragraphs and stiles
            run = prior_paragraph.add_run('Number of tests: ' + str(len(list_of_values_from_dictionary[counter_for_list_of_values_from_dictionary])) )
            run.bold = True
            run.italic = True

            # Add a new line to document
            prior_paragraph.add_run("\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run('Number of positives tests: ' + str(number_of_OK_tests[counter_for_list_of_values_from_dictionary] ))
            run.bold = True
            run.italic = True

            # Add a new line to document
            prior_paragraph.add_run("\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run('Number of negatives tests: ' + str(number_of_NOK_tests[counter_for_list_of_values_from_dictionary]))
            run.bold = True
            run.italic = True

            # Add a new line to document
            prior_paragraph.add_run("\n")

            # Add paragraphs and stiles
            run = prior_paragraph.add_run('Number of not run tests: '
                            + str(len(list_of_values_from_dictionary[counter_for_list_of_values_from_dictionary])
                                  - number_of_OK_tests[counter_for_list_of_values_from_dictionary]
                                  - number_of_NOK_tests[counter_for_list_of_values_from_dictionary]))
            run.bold = True
            run.italic = True

            counter_of_all_tests += len(list_of_values_from_dictionary[counter_for_list_of_values_from_dictionary])
            counter_for_list_of_values_from_dictionary += 1

    # ADD TOTAL STATISTICS TO DOCUMENT
    # Iterare and index all paragraphs
    for idx, paragraph in enumerate(paragraphs):
        # if found a specific paragraph
        if 'Total number of tests:' in paragraph.text:
            # Add new text to it
            run = paragraph.add_run(str(counter_of_all_tests))
            run.bold = True
            run.italic = True

        if 'Total number of positives tests:' in paragraph.text:
            # Add new text to it
            run = paragraph.add_run(str(number_of_total_OK_tests))
            run.bold = True
            run.italic = True

        if 'Total number of negatives tests:' in paragraph.text:
            # Add new text to it
            run = paragraph.add_run(str(number_of_total_NOK_tests))
            run.bold = True
            run.italic = True

        if 'Total number of not run tests:' in paragraph.text:
            # Add new text to it
            run = paragraph.add_run(str(counter_of_all_tests - number_of_total_OK_tests - number_of_total_NOK_tests))
            run.bold = True
            run.italic = True

    document.save(path_to_generated_docx_file)
    try:
        update_toc(path_to_generated_docx_file)
    except:
        print("Table of content not updated!!!")


if __name__ == "__main__":
    main()